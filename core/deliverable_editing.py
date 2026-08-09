"""Safe in-app editing primitives for workspace deliverables.

The module intentionally contains no Qt dependencies.  File inspection,
conversion, conflict detection, backup, and atomic replacement can therefore
run in background workers and be tested without constructing the desktop UI.
"""

from __future__ import annotations

import codecs
import csv
import hashlib
import io
import json
import os
import posixpath
import re
import shutil
import tempfile
import uuid
import zipfile
from dataclasses import dataclass, field
from datetime import date, datetime, time
from pathlib import Path
from typing import Any, Callable
from xml.dom import minidom
from xml.etree import ElementTree
from xml.parsers.expat import ExpatError

from core.env_utils import get_app_data_dir
from core.file_capabilities import (
    FILE_CAPABILITIES,
    OFFICE_FILE_MAX_BYTES,
    TEXT_FILE_MAX_BYTES,
    editable_extensions,
    editor_extensions,
)


MIB = 1024 * 1024
OFFICE_MAX_BYTES = OFFICE_FILE_MAX_BYTES
TEXT_MAX_BYTES = TEXT_FILE_MAX_BYTES
MAX_WORKSHEETS = 50
MAX_POPULATED_CELLS = 250_000
MAX_WORKSHEET_ROWS = 1_048_576
MAX_WORKSHEET_COLUMNS = 16_384
MAX_MERGED_RANGES = 10_000
MAX_DOCX_IMAGES = 200
MAX_OFFICE_UNCOMPRESSED_BYTES = 200 * MIB
PAYLOAD_CHUNK_CHARS = 256 * 1024

TEXT_EXTENSIONS = editor_extensions("text")
HTML_EXTENSIONS = editor_extensions("html")
TABULAR_EXTENSIONS = tuple(
    extension
    for extension, capability in FILE_CAPABILITIES.items()
    if capability.preview_kind == "table" and capability.editor_kind == "sheet"
)
DOCX_EXTENSIONS = editor_extensions("docx")
XLSX_EXTENSIONS = tuple(
    extension
    for extension, capability in FILE_CAPABILITIES.items()
    if capability.preview_kind == "xlsx" and capability.editor_kind == "sheet"
)
EDITABLE_EXTENSIONS = editable_extensions()


class DeliverableEditError(RuntimeError):
    """A user-actionable editing failure with a stable diagnostic code."""

    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = str(code or "edit_failed")
        self.message = str(message or "文件编辑失败。")


class ExternalModificationError(DeliverableEditError):
    def __init__(self, message: str = "文件已被其他程序修改，不能直接覆盖。"):
        super().__init__("external_modification", message)


@dataclass(frozen=True)
class EditorDescriptor:
    kind: str
    label: str
    extensions: tuple[str, ...]
    max_bytes: int
    web_based: bool
    visual: bool = False


@dataclass(frozen=True)
class CompatibilityIssue:
    code: str
    message: str
    blocking: bool = True
    detail: str = ""


@dataclass
class CompatibilityReport:
    path: str
    descriptor: EditorDescriptor | None
    allowed: bool
    issues: list[CompatibilityIssue] = field(default_factory=list)
    encoding: str = ""
    newline: str = "\n"
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def blocking_issues(self) -> list[CompatibilityIssue]:
        return [issue for issue in self.issues if issue.blocking]

    @property
    def message(self) -> str:
        if not self.issues:
            return ""
        return "\n".join(f"• {issue.message}" for issue in self.issues)


@dataclass
class EditSession:
    path: str
    descriptor: EditorDescriptor
    initial_fingerprint: str
    initial_size: int
    encoding: str = ""
    newline: str = "\n"
    bom_hex: str = ""
    state: str = "ready"
    dirty: bool = False
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class SaveResult:
    path: str
    backup_path: str
    fingerprint: str
    bytes_written: int


def _editor_descriptors_from_capabilities() -> tuple[EditorDescriptor, ...]:
    groups: dict[tuple[str, str, int, bool, bool], list[str]] = {}
    for extension, capability in FILE_CAPABILITIES.items():
        if not capability.editable:
            continue
        key = (
            capability.editor_kind,
            capability.editor_label,
            capability.max_bytes,
            capability.web_based,
            capability.visual,
        )
        groups.setdefault(key, []).append(extension)
    return tuple(
        EditorDescriptor(
            kind=kind,
            label=label,
            extensions=tuple(extensions),
            max_bytes=max_bytes,
            web_based=web_based,
            visual=visual,
        )
        for (kind, label, max_bytes, web_based, visual), extensions in groups.items()
    )


EDITOR_DESCRIPTORS = _editor_descriptors_from_capabilities()
EDITOR_BY_EXTENSION = {
    extension: descriptor
    for descriptor in EDITOR_DESCRIPTORS
    for extension in descriptor.extensions
}


def editor_descriptor(path_or_extension: str) -> EditorDescriptor | None:
    raw = str(path_or_extension or "").strip()
    extension = raw.lower() if raw.startswith(".") else os.path.splitext(raw)[1].lower()
    return EDITOR_BY_EXTENSION.get(extension)


def sha256_file(path: str, chunk_size: int = MIB) -> str:
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        while True:
            chunk = handle.read(chunk_size)
            if not chunk:
                break
            digest.update(chunk)
    return digest.hexdigest()


def _newline_for_text(value: str) -> str:
    crlf = value.count("\r\n")
    without_crlf = value.replace("\r\n", "")
    lf = without_crlf.count("\n")
    cr = without_crlf.count("\r")
    if crlf >= lf and crlf >= cr and crlf:
        return "\r\n"
    if cr > lf and cr:
        return "\r"
    return "\n"


def decode_text_bytes(data: bytes, selected_encoding: str = "") -> tuple[str, str, bytes, str]:
    """Decode without guessing an ambiguous legacy encoding.

    Returns text, codec name, BOM bytes, and dominant newline.  When no Unicode
    encoding can be established, the caller must ask the user to choose one.
    """

    candidates: list[tuple[bytes, str]] = [
        (codecs.BOM_UTF8, "utf-8"),
        (codecs.BOM_UTF32_LE, "utf-32-le"),
        (codecs.BOM_UTF32_BE, "utf-32-be"),
        (codecs.BOM_UTF16_LE, "utf-16-le"),
        (codecs.BOM_UTF16_BE, "utf-16-be"),
    ]
    selected = str(selected_encoding or "").strip().lower()
    if selected:
        try:
            text = data.decode(selected, errors="strict")
        except (LookupError, UnicodeDecodeError) as exc:
            raise DeliverableEditError(
                "text_decode_failed",
                f"无法使用编码 {selected_encoding} 读取文件：{exc}",
            ) from exc
        return text, selected, b"", _newline_for_text(text)

    for bom, encoding in candidates:
        if bom and data.startswith(bom):
            try:
                text = data[len(bom) :].decode(encoding, errors="strict")
            except UnicodeDecodeError as exc:
                raise DeliverableEditError(
                    "text_decode_failed",
                    f"文件声明为 {encoding}，但内容无法严格解码：{exc}",
                ) from exc
            return text, encoding, bom, _newline_for_text(text)
    try:
        text = data.decode("utf-8", errors="strict")
    except UnicodeDecodeError as exc:
        raise DeliverableEditError(
            "encoding_required",
            "无法确认文本编码。请选择编码后再进入编辑。",
        ) from exc
    return text, "utf-8", b"", _newline_for_text(text)


def encode_text_content(content: str, session: EditSession) -> bytes:
    normalized = str(content or "").replace("\r\n", "\n").replace("\r", "\n")
    newline = session.newline if session.newline in {"\n", "\r\n", "\r"} else "\n"
    normalized = normalized.replace("\n", newline)
    try:
        encoded = normalized.encode(session.encoding or "utf-8", errors="strict")
    except (LookupError, UnicodeEncodeError) as exc:
        raise DeliverableEditError(
            "text_encode_failed",
            f"内容无法使用原编码 {session.encoding or 'utf-8'} 保存：{exc}",
        ) from exc
    bom = bytes.fromhex(session.bom_hex) if session.bom_hex else b""
    return bom + encoded


def _issue(code: str, message: str, detail: str = "", blocking: bool = True) -> CompatibilityIssue:
    return CompatibilityIssue(code=code, message=message, detail=detail, blocking=blocking)


def _require_regular_file(path: str) -> tuple[str, os.stat_result]:
    normalized = os.path.abspath(os.path.normpath(str(path or "")))
    if not normalized or not os.path.isfile(normalized):
        raise DeliverableEditError("file_not_found", "文件不存在或不是普通文件。")
    return normalized, os.stat(normalized)


def _preflight_text(
    path: str,
    descriptor: EditorDescriptor,
    selected_encoding: str,
) -> CompatibilityReport:
    raw = Path(path).read_bytes()
    try:
        text, encoding, bom, newline = decode_text_bytes(raw, selected_encoding)
    except DeliverableEditError as exc:
        return CompatibilityReport(
            path=path,
            descriptor=descriptor,
            allowed=False,
            issues=[_issue(exc.code, exc.message)],
        )
    metadata: dict[str, Any] = {"bom_hex": bom.hex(), "characters": len(text)}
    if os.path.splitext(path)[1].lower() in TABULAR_EXTENSIONS:
        delimiter = "\t" if path.lower().endswith(".tsv") else ","
        try:
            rows = list(csv.reader(io.StringIO(text, newline=""), delimiter=delimiter))
        except csv.Error as exc:
            return CompatibilityReport(
                path=path,
                descriptor=descriptor,
                allowed=False,
                issues=[_issue("tabular_parse_failed", f"表格文本无法解析：{exc}")],
                encoding=encoding,
                newline=newline,
                metadata=metadata,
            )
        populated = sum(1 for row in rows for cell in row if cell != "")
        if populated > MAX_POPULATED_CELLS:
            return CompatibilityReport(
                path=path,
                descriptor=descriptor,
                allowed=False,
                issues=[
                    _issue(
                        "too_many_cells",
                        f"文件包含 {populated:,} 个非空单元格，超过内嵌编辑上限 {MAX_POPULATED_CELLS:,}。",
                    )
                ],
                encoding=encoding,
                newline=newline,
                metadata=metadata,
            )
        metadata.update(
            delimiter=delimiter,
            row_count=len(rows),
            populated_cells=populated,
        )
    return CompatibilityReport(
        path=path,
        descriptor=descriptor,
        allowed=True,
        encoding=encoding,
        newline=newline,
        metadata=metadata,
    )


_DOCX_BLOCKED_PART_PREFIXES = {
    "word/comments",
    "word/footnotes",
    "word/endnotes",
    "word/embeddings/",
    "word/activeX/",
    "word/charts/",
    "word/diagrams/",
    "word/glossary/",
    "_xmlsignatures/",
}
_DOCX_BLOCKED_XML_MARKERS = {
    b"<w:ins": "修订插入",
    b"<w:del": "修订删除",
    b"<w:sdt": "内容控件",
    b"<w:fldSimple": "复杂域",
    b"<w:instrText": "域指令",
    b"<m:oMath": "Office 公式",
    b"<w:altChunk": "外部嵌入内容",
    b"<w:object": "嵌入对象",
    b"<wp:anchor": "浮动绘图对象",
}

_DOCX_EXPORT_SCAFFOLD_PARTS = {
    "word/comments.xml",
    "word/footnotes.xml",
}
_DOCX_EXPORT_SCAFFOLD_REFERENCE_MARKERS = (
    b"<w:commentRangeStart",
    b"<w:commentRangeEnd",
    b"<w:commentReference",
    b"<w:footnoteReference",
)

_DOCX_WORDPROCESSING_NS = (
    "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
)
_DOCX_OFFICE_REL_NS = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
)
_DOCX_PACKAGE_REL_NS = (
    "http://schemas.openxmlformats.org/package/2006/relationships"
)
_DOCX_CONTENT_TYPES_NS = (
    "http://schemas.openxmlformats.org/package/2006/content-types"
)
_DOCX_HEADER_REL_TYPE = f"{_DOCX_OFFICE_REL_NS}/header"
_DOCX_FOOTER_REL_TYPE = f"{_DOCX_OFFICE_REL_NS}/footer"
_DOCX_HEADER_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"
)
_DOCX_FOOTER_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"
)


def _is_safe_docx_export_scaffold(
    archive: zipfile.ZipFile,
    part_name: str,
    document_xml: bytes,
) -> bool:
    """Accept only the empty OOXML parts always emitted by the DOCX exporter."""
    normalized_name = part_name.replace("\\", "/").lower()
    if normalized_name not in _DOCX_EXPORT_SCAFFOLD_PARTS:
        return False
    if any(marker in document_xml for marker in _DOCX_EXPORT_SCAFFOLD_REFERENCE_MARKERS):
        return False
    try:
        root = ElementTree.fromstring(archive.read(part_name))
    except (KeyError, ElementTree.ParseError):
        return False
    namespace = f"{{{_DOCX_WORDPROCESSING_NS}}}"
    if normalized_name == "word/comments.xml":
        return root.tag == f"{namespace}comments" and not list(root)
    if root.tag != f"{namespace}footnotes":
        return False
    allowed_descendants = {
        f"{namespace}footnote",
        f"{namespace}p",
        f"{namespace}pPr",
        f"{namespace}spacing",
        f"{namespace}r",
        f"{namespace}rPr",
        f"{namespace}rStyle",
        f"{namespace}footnoteRef",
        f"{namespace}separator",
        f"{namespace}continuationSeparator",
    }
    id_attribute = f"{namespace}id"
    type_attribute = f"{namespace}type"
    expected_types = {"-1": "separator", "0": "continuationSeparator"}
    footnotes = list(root)
    if len(footnotes) != len(expected_types):
        return False
    seen_ids: set[str] = set()
    for footnote in footnotes:
        footnote_id = str(footnote.attrib.get(id_attribute) or "")
        if (
            footnote.tag != f"{namespace}footnote"
            or footnote_id in seen_ids
            or footnote.attrib.get(type_attribute) != expected_types.get(footnote_id)
            or any(node.tag not in allowed_descendants for node in footnote.iter())
        ):
            return False
        seen_ids.add(footnote_id)
    return seen_ids == set(expected_types)


def _docx_part_relationships_name(part_name: str) -> str:
    directory, filename = posixpath.split(part_name)
    return posixpath.join(directory, "_rels", f"{filename}.rels")


def _resolve_docx_relationship_target(source_part: str, target: str) -> str:
    normalized_target = str(target or "").replace("\\", "/")
    if not normalized_target:
        raise DeliverableEditError(
            "invalid_docx",
            "DOCX 页眉或页脚关系缺少目标部件。",
        )
    if normalized_target.startswith("/"):
        normalized = posixpath.normpath(normalized_target.lstrip("/"))
    else:
        normalized = posixpath.normpath(
            posixpath.join(posixpath.dirname(source_part), normalized_target)
        )
    if normalized in {"", ".", ".."} or normalized.startswith("../"):
        raise DeliverableEditError(
            "invalid_docx",
            "DOCX 页眉或页脚关系指向包外路径。",
        )
    return normalized


def _inspect_docx_header_footer(archive, names: set[str]) -> tuple[dict[str, Any] | None, list[CompatibilityIssue]]:
    document_root = ElementTree.fromstring(archive.read("word/document.xml"))
    section_tag = f"{{{_DOCX_WORDPROCESSING_NS}}}sectPr"
    header_tag = f"{{{_DOCX_WORDPROCESSING_NS}}}headerReference"
    footer_tag = f"{{{_DOCX_WORDPROCESSING_NS}}}footerReference"
    relationship_id_attr = f"{{{_DOCX_OFFICE_REL_NS}}}id"
    section_nodes = document_root.findall(f".//{section_tag}")
    references = [
        child
        for section in section_nodes
        for child in list(section)
        if child.tag in {header_tag, footer_tag}
    ]
    if not references:
        return None, []
    if len(section_nodes) != 1:
        return None, [
            _issue(
                "docx_complex_header_footer",
                "文档包含多节页眉或页脚，当前编辑器无法保证完整往返。",
            )
        ]

    relationships_name = "word/_rels/document.xml.rels"
    if relationships_name not in names:
        raise DeliverableEditError(
            "invalid_docx",
            "DOCX 缺少正文关系表，无法解析页眉或页脚。",
        )
    relationships_root = ElementTree.fromstring(archive.read(relationships_name))
    relationship_tag = f"{{{_DOCX_PACKAGE_REL_NS}}}Relationship"
    relationship_by_id = {
        str(node.attrib.get("Id") or ""): node
        for node in relationships_root.findall(relationship_tag)
    }
    plan_relationships: dict[str, dict[str, str]] = {}
    preserved_parts: dict[str, bytes] = {}
    issues: list[CompatibilityIssue] = []
    for reference in references:
        reference_type = str(
            reference.attrib.get(f"{{{_DOCX_WORDPROCESSING_NS}}}type") or "default"
        )
        if reference_type != "default":
            issues.append(
                _issue(
                    "docx_complex_header_footer",
                    "文档使用首页或奇偶页专用页眉页脚，当前编辑器无法保证完整往返。",
                )
            )
            break
        relationship_id = str(reference.attrib.get(relationship_id_attr) or "")
        relationship = relationship_by_id.get(relationship_id)
        kind = "header" if reference.tag == header_tag else "footer"
        expected_type = (
            _DOCX_HEADER_REL_TYPE if kind == "header" else _DOCX_FOOTER_REL_TYPE
        )
        if (
            relationship is None
            or str(relationship.attrib.get("Type") or "") != expected_type
            or str(relationship.attrib.get("TargetMode") or "").lower() == "external"
        ):
            raise DeliverableEditError(
                "invalid_docx",
                "DOCX 页眉或页脚关系无效。",
            )
        part_name = _resolve_docx_relationship_target(
            "word/document.xml",
            str(relationship.attrib.get("Target") or ""),
        )
        if part_name not in names:
            raise DeliverableEditError(
                "invalid_docx",
                f"DOCX 缺少页眉或页脚部件：{part_name}",
            )
        related_relationships_name = _docx_part_relationships_name(part_name)
        if related_relationships_name in names:
            related_root = ElementTree.fromstring(
                archive.read(related_relationships_name)
            )
            if related_root.findall(relationship_tag):
                issues.append(
                    _issue(
                        "docx_complex_header_footer",
                        "页眉或页脚包含图片、链接等关联资源，当前编辑器无法保证完整往返。",
                    )
                )
                break
        part_bytes = archive.read(part_name)
        part_root = ElementTree.fromstring(part_bytes)
        if any(
            attribute.startswith(f"{{{_DOCX_OFFICE_REL_NS}}}")
            for node in part_root.iter()
            for attribute in node.attrib
        ):
            issues.append(
                _issue(
                    "docx_complex_header_footer",
                    "页眉或页脚包含外部关联内容，当前编辑器无法保证完整往返。",
                )
            )
            break
        plan_relationships[relationship_id] = {
            "kind": kind,
            "part_name": part_name,
        }
        preserved_parts[part_name] = part_bytes

    if issues:
        return None, issues
    return {
        "section_xml": ElementTree.tostring(section_nodes[0], encoding="utf-8"),
        "relationships": plan_relationships,
        "parts": preserved_parts,
    }, []


def _preflight_docx(
    path: str,
    descriptor: EditorDescriptor,
) -> CompatibilityReport:
    issues: list[CompatibilityIssue] = []
    metadata: dict[str, Any] = {}
    try:
        with zipfile.ZipFile(path) as archive:
            names = {name.replace("\\", "/") for name in archive.namelist()}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise DeliverableEditError("invalid_docx", "文件不是有效的 DOCX OOXML 文档。")
            document_xml = archive.read("word/document.xml")
            for name in sorted(names):
                lowered = name.lower()
                if lowered == "word/vbaproject.bin":
                    issues.append(_issue("docx_macro", "文档包含宏，不能在应用内安全编辑。"))
                for prefix in _DOCX_BLOCKED_PART_PREFIXES:
                    if lowered.startswith(prefix.lower()):
                        if _is_safe_docx_export_scaffold(
                            archive,
                            name,
                            document_xml,
                        ):
                            break
                        issues.append(
                            _issue(
                                "docx_unsupported_part",
                                f"文档包含不支持的结构：{name}",
                            )
                        )
                        break
            for marker, label in _DOCX_BLOCKED_XML_MARKERS.items():
                if marker in document_xml:
                    issues.append(
                        _issue(
                            "docx_unsupported_markup",
                            f"文档包含{label}，当前编辑器无法保证完整往返。",
                        )
                    )
            media_bytes = sum(
                archive.getinfo(name).file_size
                for name in names
                if name.lower().startswith("word/media/")
            )
            image_count = sum(
                1 for name in names if name.lower().startswith("word/media/")
            )
            uncompressed_bytes = sum(item.file_size for item in archive.infolist())
            metadata["media_bytes"] = media_bytes
            metadata["image_count"] = image_count
            metadata["uncompressed_bytes"] = uncompressed_bytes
            if media_bytes > 20 * MIB:
                issues.append(
                    _issue(
                        "docx_media_too_large",
                        "文档内嵌媒体超过 20 MiB，可能导致编辑器内存占用过高。",
                    )
                )
            if image_count > MAX_DOCX_IMAGES:
                issues.append(
                    _issue(
                        "docx_too_many_images",
                        f"文档包含 {image_count} 个图片资源，超过上限 {MAX_DOCX_IMAGES}。",
                    )
                )
            if uncompressed_bytes > MAX_OFFICE_UNCOMPRESSED_BYTES:
                issues.append(
                    _issue(
                        "office_uncompressed_too_large",
                        "DOCX 解压后内容超过 200 MiB，超出安全内存预算。",
                    )
                )
            header_footer_plan, header_footer_issues = _inspect_docx_header_footer(
                archive,
                names,
            )
            issues.extend(header_footer_issues)
            if header_footer_plan is not None:
                metadata["_docx_header_footer_plan"] = header_footer_plan
                metadata["docx_preserved_header_footer_count"] = len(
                    header_footer_plan["relationships"]
                )
    except zipfile.BadZipFile as exc:
        raise DeliverableEditError("invalid_docx", "文件不是有效的 DOCX 压缩包。") from exc
    except ElementTree.ParseError as exc:
        raise DeliverableEditError("invalid_docx", f"DOCX XML 结构无效：{exc}") from exc

    if not issues:
        try:
            from docx import Document

            document = Document(path)
            metadata["paragraphs"] = len(document.paragraphs)
            metadata["tables"] = len(document.tables)
        except Exception as exc:
            issues.append(_issue("docx_open_failed", f"DOCX 结构验证失败：{exc}"))
    return CompatibilityReport(
        path=path,
        descriptor=descriptor,
        allowed=not any(item.blocking for item in issues),
        issues=issues,
        metadata=metadata,
    )


_XLSX_BLOCKED_PREFIXES = {
    "xl/charts/",
    "xl/drawings/",
    "xl/pivotTables/",
    "xl/pivotCache/",
    "xl/slicers/",
    "xl/externalLinks/",
    "xl/embeddings/",
    "xl/activeX/",
    "xl/connections",
    "xl/queryTables/",
    "xl/tables/",
    "xl/model/",
    "_xmlsignatures/",
}


def _xlsx_populated_cell_count(archive: zipfile.ZipFile, names: set[str]) -> int:
    total = 0
    cell_suffix = "}c"
    for name in sorted(names):
        lowered = name.lower()
        if not lowered.startswith("xl/worksheets/sheet") or not lowered.endswith(".xml"):
            continue
        with archive.open(name) as stream:
            try:
                for _event, element in ElementTree.iterparse(stream, events=("end",)):
                    if element.tag.endswith(cell_suffix):
                        total += 1
                        if total > MAX_POPULATED_CELLS:
                            return total
                    element.clear()
            except ElementTree.ParseError as exc:
                raise DeliverableEditError(
                    "xlsx_xml_invalid",
                    f"工作表 XML 无法解析：{name}：{exc}",
                ) from exc
    return total


def _preflight_xlsx(path: str, descriptor: EditorDescriptor) -> CompatibilityReport:
    issues: list[CompatibilityIssue] = []
    metadata: dict[str, Any] = {}
    try:
        with zipfile.ZipFile(path) as archive:
            names = {name.replace("\\", "/") for name in archive.namelist()}
            if "[Content_Types].xml" not in names or "xl/workbook.xml" not in names:
                raise DeliverableEditError("invalid_xlsx", "文件不是有效的 XLSX OOXML 工作簿。")
            for name in sorted(names):
                lowered = name.lower()
                if lowered.endswith("vbaproject.bin"):
                    issues.append(_issue("xlsx_macro", "工作簿包含宏，不能在应用内安全编辑。"))
                for prefix in _XLSX_BLOCKED_PREFIXES:
                    if lowered.startswith(prefix.lower()):
                        issues.append(
                            _issue(
                                "xlsx_unsupported_part",
                                f"工作簿包含不支持的结构：{name}",
                            )
                        )
                        break
            workbook_xml = archive.read("xl/workbook.xml")
            try:
                workbook_root = ElementTree.fromstring(workbook_xml)
            except ElementTree.ParseError as exc:
                raise DeliverableEditError(
                    "xlsx_xml_invalid",
                    f"工作簿 XML 无法解析：{exc}",
                ) from exc
            protection = next(
                (
                    node
                    for node in workbook_root.iter()
                    if node.tag.endswith("}workbookProtection") or node.tag == "workbookProtection"
                ),
                None,
            )
            if protection is not None and protection.attrib:
                issues.append(_issue("xlsx_protected", "工作簿已启用结构保护，不能在应用内编辑。"))
            populated = _xlsx_populated_cell_count(archive, names)
            metadata["populated_cells"] = populated
            uncompressed_bytes = sum(item.file_size for item in archive.infolist())
            metadata["uncompressed_bytes"] = uncompressed_bytes
            if populated > MAX_POPULATED_CELLS:
                issues.append(
                    _issue(
                        "too_many_cells",
                        f"工作簿包含超过 {MAX_POPULATED_CELLS:,} 个有效单元格，超出内嵌编辑预算。",
                    )
                )
            if uncompressed_bytes > MAX_OFFICE_UNCOMPRESSED_BYTES:
                issues.append(
                    _issue(
                        "office_uncompressed_too_large",
                        "XLSX 解压后内容超过 200 MiB，超出安全内存预算。",
                    )
                )
    except zipfile.BadZipFile as exc:
        raise DeliverableEditError("invalid_xlsx", "文件不是有效的 XLSX 压缩包。") from exc

    if not issues:
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(
                io.BytesIO(Path(path).read_bytes()),
                read_only=True,
                data_only=False,
            )
            try:
                metadata["worksheets"] = len(workbook.sheetnames)
                if len(workbook.sheetnames) > MAX_WORKSHEETS:
                    issues.append(
                        _issue(
                            "too_many_worksheets",
                            f"工作簿包含 {len(workbook.sheetnames)} 个工作表，超过上限 {MAX_WORKSHEETS}。",
                        )
                    )
            finally:
                workbook.close()
        except Exception as exc:
            issues.append(_issue("xlsx_open_failed", f"XLSX 结构验证失败：{exc}"))
    return CompatibilityReport(
        path=path,
        descriptor=descriptor,
        allowed=not any(item.blocking for item in issues),
        issues=issues,
        metadata=metadata,
    )


def preflight_edit(path: str, selected_encoding: str = "") -> CompatibilityReport:
    normalized, stat = _require_regular_file(path)
    descriptor = editor_descriptor(normalized)
    if descriptor is None:
        return CompatibilityReport(
            path=normalized,
            descriptor=None,
            allowed=False,
            issues=[_issue("unsupported_format", "当前格式不支持应用内编辑。")],
        )
    if stat.st_size > descriptor.max_bytes:
        return CompatibilityReport(
            path=normalized,
            descriptor=descriptor,
            allowed=False,
            issues=[
                _issue(
                    "file_too_large",
                    f"文件大小 {stat.st_size / MIB:.1f} MiB，超过 {descriptor.max_bytes // MIB} MiB 编辑上限。",
                )
            ],
        )
    if descriptor.kind in {"text", "html"} or os.path.splitext(normalized)[1].lower() in TABULAR_EXTENSIONS:
        report = _preflight_text(normalized, descriptor, selected_encoding)
    elif descriptor.kind == "docx":
        report = _preflight_docx(normalized, descriptor)
    elif descriptor.kind == "sheet":
        report = _preflight_xlsx(normalized, descriptor)
    else:
        report = CompatibilityReport(
            path=normalized,
            descriptor=descriptor,
            allowed=False,
            issues=[_issue("unsupported_format", "当前格式不支持应用内编辑。")],
        )
    report.metadata.setdefault("size", stat.st_size)
    return report


def create_edit_session(path: str, selected_encoding: str = "") -> tuple[EditSession, CompatibilityReport]:
    report = preflight_edit(path, selected_encoding=selected_encoding)
    if not report.allowed or report.descriptor is None:
        raise DeliverableEditError(
            report.blocking_issues[0].code if report.blocking_issues else "preflight_failed",
            report.message or "文件未通过兼容性预检。",
        )
    normalized = report.path
    session = EditSession(
        path=normalized,
        descriptor=report.descriptor,
        initial_fingerprint=sha256_file(normalized),
        initial_size=os.path.getsize(normalized),
        encoding=report.encoding,
        newline=report.newline,
        bom_hex=str(report.metadata.get("bom_hex") or ""),
        metadata=dict(report.metadata),
    )
    return session, report


_HTML_PRESERVED_TAGS = ("script", "iframe", "object", "embed")


def prepare_html_edit_source(source: str, base_url: str = "") -> tuple[str, list[str]]:
    """Build a script-free editing copy while retaining restorable nodes."""

    from bs4 import BeautifulSoup

    soup = BeautifulSoup(str(source or ""), "html.parser")
    for node in soup.find_all(True):
        if any(str(attribute).lower().startswith("data-cowork-") for attribute in node.attrs):
            raise DeliverableEditError(
                "html_reserved_attribute",
                "HTML 使用了编辑器保留的 data-cowork-* 属性，不能安全进入编辑。",
            )
    if soup.html is None:
        wrapped = BeautifulSoup("<!doctype html><html><head></head><body></body></html>", "html.parser")
        for child in list(soup.contents):
            wrapped.body.append(child.extract())
        soup = wrapped
    if soup.head is None:
        head = soup.new_tag("head")
        soup.html.insert(0, head)
    if soup.body is None:
        body = soup.new_tag("body")
        soup.html.append(body)

    preserved: list[str] = []
    candidates = list(soup.find_all(_HTML_PRESERVED_TAGS))
    candidates.extend(
        node
        for node in soup.find_all("meta")
        if str(node.get("http-equiv") or "").strip().lower() == "refresh"
    )
    for node in candidates:
        index = len(preserved)
        preserved.append(str(node))
        placeholder = soup.new_tag("template")
        placeholder["data-cowork-preserved-node"] = str(index)
        node.replace_with(placeholder)
    if base_url:
        head = soup.head
        if head is None:
            html_node = soup.html
            if html_node is None:
                html_node = soup.new_tag("html")
                for child in list(soup.contents):
                    html_node.append(child.extract())
                soup.append(html_node)
            head = soup.new_tag("head")
            html_node.insert(0, head)
        base = soup.new_tag("base", href=str(base_url))
        base["data-cowork-edit-base"] = "true"
        head.insert(0, base)
    return str(soup), preserved


def restore_html_edit_source(edited_source: str, preserved: list[str]) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(str(edited_source or ""), "html.parser")
    placeholders = soup.find_all("template", attrs={"data-cowork-preserved-node": True})
    seen: set[int] = set()
    for placeholder in placeholders:
        raw_index = str(placeholder.get("data-cowork-preserved-node") or "")
        try:
            index = int(raw_index)
        except ValueError as exc:
            raise DeliverableEditError(
                "html_preserved_node_invalid",
                "HTML 安全占位符已损坏，不能保存。",
            ) from exc
        if index < 0 or index >= len(preserved) or index in seen:
            raise DeliverableEditError(
                "html_preserved_node_invalid",
                "HTML 安全占位符与原文件不一致，不能保存。",
            )
        seen.add(index)
        fragment = BeautifulSoup(preserved[index], "html.parser")
        for child in list(fragment.contents):
            placeholder.insert_before(child.extract())
        placeholder.extract()
    if len(seen) != len(preserved):
        raise DeliverableEditError(
            "html_preserved_node_missing",
            "编辑后的页面缺少受保护的脚本或嵌入节点，不能覆盖原文件。",
        )
    for node in list(soup.find_all("base", attrs={"data-cowork-edit-base": True})):
        node.decompose()
    for node in soup.find_all(True):
        for attribute in list(node.attrs):
            lowered = str(attribute).lower()
            if lowered in {"contenteditable", "spellcheck"} or lowered.startswith("data-cowork-edit"):
                del node.attrs[attribute]
    rendered = str(soup)
    if not re.match(r"^\s*<!doctype\b", rendered, flags=re.IGNORECASE):
        rendered = "<!doctype html>\n" + rendered
    return rendered


def apply_html_dom_patch(
    original_source: str,
    edited_source: str,
    preserved: list[str],
) -> str:
    """Apply the isolated editor's body DOM back onto the original document.

    The original head and document-level nodes remain authoritative.  Only the
    editable body attributes and children are replaced after all protected
    nodes have been proven present and restored.
    """

    from bs4 import BeautifulSoup

    restored_edit_copy = restore_html_edit_source(edited_source, preserved)
    original = BeautifulSoup(str(original_source or ""), "html.parser")
    edited = BeautifulSoup(restored_edit_copy, "html.parser")
    if original.body is None or edited.body is None:
        return restored_edit_copy

    original.body.attrs = {
        key: value
        for key, value in edited.body.attrs.items()
        if not str(key).lower().startswith("data-cowork-")
        and str(key).lower() not in {"contenteditable", "spellcheck"}
    }
    original.body.clear()
    for child in list(edited.body.contents):
        original.body.append(child.extract())
    rendered = str(original)
    if not re.match(r"^\s*<!doctype\b", rendered, flags=re.IGNORECASE):
        rendered = "<!doctype html>\n" + rendered
    return rendered


def _color_hex(color: Any) -> str:
    if color is None:
        return ""
    color_type = str(getattr(color, "type", "") or "")
    value = str(getattr(color, "rgb", "") or "")
    if color_type == "rgb" and value:
        value = value[-6:]
        if re.fullmatch(r"[0-9a-fA-F]{6}", value):
            return "#" + value.upper()
    return ""


def _openpyxl_border_style(name: str) -> int:
    mapping = {
        "thin": 1,
        "hair": 2,
        "dotted": 3,
        "dashed": 4,
        "dashDot": 5,
        "dashDotDot": 6,
        "double": 7,
        "medium": 8,
        "mediumDashed": 9,
        "mediumDashDot": 10,
        "mediumDashDotDot": 11,
        "slantDashDot": 12,
        "thick": 13,
    }
    return mapping.get(str(name or ""), 0)


def _univer_border_style(value: int) -> str | None:
    mapping = {
        1: "thin",
        2: "hair",
        3: "dotted",
        4: "dashed",
        5: "dashDot",
        6: "dashDotDot",
        7: "double",
        8: "medium",
        9: "mediumDashed",
        10: "mediumDashDot",
        11: "mediumDashDotDot",
        12: "slantDashDot",
        13: "thick",
    }
    return mapping.get(int(value or 0))


def _cell_style_to_univer(cell: Any) -> dict[str, Any]:
    style: dict[str, Any] = {}
    font = cell.font
    if font:
        if font.name:
            style["ff"] = font.name
        if font.sz:
            style["fs"] = float(font.sz)
        if font.bold:
            style["bl"] = 1
        if font.italic:
            style["it"] = 1
        if font.underline:
            style["ul"] = {"s": 1, "t": 10 if font.underline == "double" else 12}
        if font.strike:
            style["st"] = {"s": 1}
        color = _color_hex(font.color)
        if color:
            style["cl"] = {"rgb": color}
    fill = cell.fill
    if fill and fill.fill_type == "solid":
        fill_color = _color_hex(fill.fgColor)
        if fill_color:
            style["bg"] = {"rgb": fill_color}
    alignment = cell.alignment
    if alignment:
        horizontal = {
            "left": 1,
            "center": 2,
            "right": 3,
            "justify": 4,
            "distributed": 6,
        }.get(alignment.horizontal)
        vertical = {"top": 1, "center": 2, "bottom": 3}.get(alignment.vertical)
        if horizontal:
            style["ht"] = horizontal
        if vertical:
            style["vt"] = vertical
        if alignment.wrap_text:
            style["tb"] = 3
        elif alignment.shrink_to_fit:
            style["tb"] = 2
        if alignment.text_rotation:
            style["tr"] = {"a": int(alignment.text_rotation)}
    if cell.number_format and cell.number_format != "General":
        style["n"] = {"pattern": cell.number_format}
    border_data: dict[str, Any] = {}
    for key, side in (
        ("t", cell.border.top),
        ("r", cell.border.right),
        ("b", cell.border.bottom),
        ("l", cell.border.left),
    ):
        style_type = _openpyxl_border_style(getattr(side, "style", ""))
        if not style_type:
            continue
        color = _color_hex(getattr(side, "color", None)) or "#000000"
        border_data[key] = {"s": style_type, "cl": {"rgb": color}}
    if border_data:
        style["bd"] = border_data
    return style


def _cell_value_to_univer(cell: Any) -> dict[str, Any]:
    value = cell.value
    payload: dict[str, Any] = {}
    if cell.data_type == "f":
        formula = str(value or "")
        payload["f"] = formula if formula.startswith("=") else "=" + formula
        return payload
    if value is None:
        return payload
    if isinstance(value, bool):
        payload.update(v=value, t=3)
    elif isinstance(value, (int, float)):
        payload.update(v=value, t=2)
    elif isinstance(value, (datetime, date, time)):
        from openpyxl.utils.datetime import to_excel

        payload.update(v=float(to_excel(value)), t=2)
    else:
        payload.update(v=str(value), t=1)
    return payload


def _safe_sheet_id(index: int, title: str) -> str:
    digest = hashlib.sha1(f"{index}:{title}".encode("utf-8")).hexdigest()[:12]
    return f"sheet-{digest}"


def xlsx_to_univer_snapshot(path: str) -> dict[str, Any]:
    import openpyxl
    from openpyxl.utils.cell import coordinate_to_tuple

    workbook = openpyxl.load_workbook(path, read_only=False, data_only=False)
    styles: dict[str, Any] = {}
    sheets: dict[str, Any] = {}
    order: list[str] = []
    try:
        for sheet_index, worksheet in enumerate(workbook.worksheets):
            sheet_id = _safe_sheet_id(sheet_index, worksheet.title)
            order.append(sheet_id)
            frozen_row = 0
            frozen_column = 0
            if worksheet.freeze_panes:
                frozen_row, frozen_column = coordinate_to_tuple(
                    str(worksheet.freeze_panes)
                )
                frozen_row = max(0, frozen_row - 1)
                frozen_column = max(0, frozen_column - 1)
            cell_data: dict[str, dict[str, Any]] = {}
            for cell in sorted(
                worksheet._cells.values(),  # noqa: SLF001 - avoids materialising empty ranges
                key=lambda item: (item.row, item.column),
            ):
                payload = _cell_value_to_univer(cell)
                if cell.has_style:
                    style_id = f"s{cell.style_id}"
                    styles.setdefault(style_id, _cell_style_to_univer(cell))
                    payload["s"] = style_id
                if payload:
                    cell_data.setdefault(str(cell.row - 1), {})[str(cell.column - 1)] = payload
            row_data: dict[str, Any] = {}
            for index, dimension in worksheet.row_dimensions.items():
                entry: dict[str, Any] = {}
                if dimension.height is not None:
                    entry["h"] = round(float(dimension.height) / 0.75, 2)
                if dimension.hidden:
                    entry["hd"] = 1
                if entry:
                    row_data[str(int(index) - 1)] = entry
            column_data: dict[str, Any] = {}
            for key, dimension in worksheet.column_dimensions.items():
                from openpyxl.utils.cell import column_index_from_string

                try:
                    index = column_index_from_string(key) - 1
                except ValueError:
                    continue
                entry = {}
                if dimension.width is not None:
                    entry["w"] = round(float(dimension.width) * 7 + 5, 2)
                if dimension.hidden:
                    entry["hd"] = 1
                if entry:
                    column_data[str(index)] = entry
            merge_data = [
                {
                    "startRow": merged.min_row - 1,
                    "startColumn": merged.min_col - 1,
                    "endRow": merged.max_row - 1,
                    "endColumn": merged.max_col - 1,
                }
                for merged in worksheet.merged_cells.ranges
            ]
            sheets[sheet_id] = {
                "id": sheet_id,
                "name": worksheet.title,
                "tabColor": _color_hex(worksheet.sheet_properties.tabColor),
                "hidden": 1 if worksheet.sheet_state != "visible" else 0,
                "rowCount": max(100, min(max(worksheet.max_row + 20, 100), 1_000_000)),
                "columnCount": max(26, min(max(worksheet.max_column + 10, 26), 16_384)),
                "freeze": {
                    "xSplit": frozen_column,
                    "ySplit": frozen_row,
                    "startRow": frozen_row,
                    "startColumn": frozen_column,
                },
                "zoomRatio": 1,
                "scrollTop": 0,
                "scrollLeft": 0,
                "defaultColumnWidth": 88,
                "defaultRowHeight": 20,
                "mergeData": merge_data,
                "cellData": cell_data,
                "rowData": row_data,
                "columnData": column_data,
                "rowHeader": {"width": 46},
                "columnHeader": {"height": 20},
                "showGridlines": 1 if worksheet.sheet_view.showGridLines is not False else 0,
                "rightToLeft": 1 if worksheet.sheet_view.rightToLeft else 0,
            }
    finally:
        workbook.close()
    return {
        "id": f"workbook-{hashlib.sha1(os.path.abspath(path).encode('utf-8')).hexdigest()[:12]}",
        "name": os.path.basename(path),
        "appVersion": "0.25.1",
        "locale": "zhCN",
        "styles": styles,
        "sheetOrder": order,
        "sheets": sheets,
    }


def rows_to_univer_snapshot(rows: list[list[str]], name: str = "数据") -> dict[str, Any]:
    sheet_id = _safe_sheet_id(0, name)
    cell_data: dict[str, dict[str, Any]] = {}
    max_columns = 0
    for row_index, row in enumerate(rows):
        max_columns = max(max_columns, len(row))
        for column_index, value in enumerate(row):
            if value == "":
                continue
            cell_data.setdefault(str(row_index), {})[str(column_index)] = {"v": str(value), "t": 1}
    return {
        "id": f"workbook-{uuid.uuid4().hex[:12]}",
        "name": name,
        "appVersion": "0.25.1",
        "locale": "zhCN",
        "styles": {},
        "sheetOrder": [sheet_id],
        "sheets": {
            sheet_id: {
                "id": sheet_id,
                "name": name,
                "hidden": 0,
                "rowCount": max(100, len(rows) + 20),
                "columnCount": max(26, max_columns + 10),
                "freeze": {"xSplit": 0, "ySplit": 0, "startRow": 0, "startColumn": 0},
                "zoomRatio": 1,
                "scrollTop": 0,
                "scrollLeft": 0,
                "defaultColumnWidth": 88,
                "defaultRowHeight": 20,
                "mergeData": [],
                "cellData": cell_data,
                "rowData": {},
                "columnData": {},
                "rowHeader": {"width": 46},
                "columnHeader": {"height": 20},
                "showGridlines": 1,
                "rightToLeft": 0,
            }
        },
    }


def _univer_color(value: Any) -> str:
    if isinstance(value, dict):
        raw = str(value.get("rgb") or "")
        if re.fullmatch(r"#[0-9A-Fa-f]{6}", raw):
            return "FF" + raw[1:].upper()
    return ""


def _apply_univer_style(cell: Any, style: dict[str, Any]) -> None:
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    if not isinstance(style, dict):
        return
    underline = style.get("ul") if isinstance(style.get("ul"), dict) else {}
    strike = style.get("st") if isinstance(style.get("st"), dict) else {}
    font_color = _univer_color(style.get("cl"))
    cell.font = Font(
        name=str(style.get("ff") or "Calibri"),
        size=float(style.get("fs") or 11),
        bold=bool(style.get("bl")),
        italic=bool(style.get("it")),
        underline="double" if underline.get("t") == 10 else ("single" if underline.get("s") else None),
        strike=bool(strike.get("s")),
        color=font_color or None,
    )
    fill_color = _univer_color(style.get("bg"))
    if fill_color:
        cell.fill = PatternFill(fill_type="solid", fgColor=fill_color)
    horizontal = {1: "left", 2: "center", 3: "right", 4: "justify", 6: "distributed"}.get(
        int(style.get("ht") or 0)
    )
    vertical = {1: "top", 2: "center", 3: "bottom"}.get(int(style.get("vt") or 0))
    rotation = style.get("tr") if isinstance(style.get("tr"), dict) else {}
    cell.alignment = Alignment(
        horizontal=horizontal,
        vertical=vertical,
        wrap_text=int(style.get("tb") or 0) == 3,
        shrink_to_fit=int(style.get("tb") or 0) == 2,
        text_rotation=max(0, min(180, int(rotation.get("a") or 0))),
    )
    number_format = style.get("n")
    if isinstance(number_format, dict) and number_format.get("pattern"):
        cell.number_format = str(number_format["pattern"])
    border = style.get("bd")
    if isinstance(border, dict):
        sides: dict[str, Side] = {}
        for key, openpyxl_key in (("t", "top"), ("r", "right"), ("b", "bottom"), ("l", "left")):
            value = border.get(key)
            if not isinstance(value, dict):
                continue
            border_style = _univer_border_style(int(value.get("s") or 0))
            if border_style:
                sides[openpyxl_key] = Side(
                    style=border_style,
                    color=_univer_color(value.get("cl")) or "FF000000",
                )
        if sides:
            cell.border = Border(**sides)


def _validate_snapshot(snapshot: dict[str, Any]) -> tuple[list[str], dict[str, Any]]:
    if not isinstance(snapshot, dict):
        raise DeliverableEditError("sheet_payload_invalid", "表格编辑器返回了无效数据。")
    order = snapshot.get("sheetOrder")
    sheets = snapshot.get("sheets")
    if not isinstance(order, list) or not isinstance(sheets, dict) or not order:
        raise DeliverableEditError("sheet_payload_invalid", "工作簿必须至少包含一个工作表。")
    normalized_order = [str(item) for item in order]
    if len(normalized_order) != len(set(normalized_order)):
        raise DeliverableEditError("sheet_payload_invalid", "工作表顺序包含重复项。")
    if len(normalized_order) > MAX_WORKSHEETS:
        raise DeliverableEditError(
            "too_many_worksheets",
            f"工作表数量超过上限 {MAX_WORKSHEETS}。",
        )
    extra_sheet_ids = set(str(item) for item in sheets) - set(normalized_order)
    if extra_sheet_ids:
        raise DeliverableEditError(
            "sheet_payload_invalid",
            "工作簿包含未列入顺序的工作表，不能在保存时静默丢弃。",
        )
    populated = 0
    sheet_names: set[str] = set()
    for sheet_id in normalized_order:
        sheet = sheets.get(sheet_id)
        if not isinstance(sheet, dict):
            raise DeliverableEditError("sheet_payload_invalid", "工作表顺序与数据不一致。")
        sheet_name = str(sheet.get("name") or "").strip()
        if (
            not sheet_name
            or len(sheet_name) > 31
            or re.search(r"[\[\]:*?/\\]", sheet_name)
        ):
            raise DeliverableEditError(
                "sheet_name_invalid",
                f"工作表名称“{sheet_name or '(空)'}”不符合 XLSX 规则。",
            )
        normalized_name = sheet_name.casefold()
        if normalized_name in sheet_names:
            raise DeliverableEditError(
                "sheet_name_duplicate",
                f"工作表名称“{sheet_name}”重复。",
            )
        sheet_names.add(normalized_name)
        cell_data = sheet.get("cellData") or {}
        if not isinstance(cell_data, dict):
            raise DeliverableEditError("sheet_payload_invalid", "单元格数据结构无效。")
        for row_key, row in cell_data.items():
            try:
                row_index = int(row_key)
            except (TypeError, ValueError) as exc:
                raise DeliverableEditError(
                    "sheet_payload_invalid",
                    "工作表包含无效的行索引。",
                ) from exc
            if row_index < 0 or row_index >= MAX_WORKSHEET_ROWS:
                raise DeliverableEditError(
                    "sheet_bounds_invalid",
                    f"行索引超出 XLSX 上限 {MAX_WORKSHEET_ROWS:,}。",
                )
            if not isinstance(row, dict):
                raise DeliverableEditError(
                    "sheet_payload_invalid",
                    "工作表包含无效的单元格行数据。",
                )
            for column_key, value in row.items():
                try:
                    column_index = int(column_key)
                except (TypeError, ValueError) as exc:
                    raise DeliverableEditError(
                        "sheet_payload_invalid",
                        "工作表包含无效的列索引。",
                    ) from exc
                if column_index < 0 or column_index >= MAX_WORKSHEET_COLUMNS:
                    raise DeliverableEditError(
                        "sheet_bounds_invalid",
                        f"列索引超出 XLSX 上限 {MAX_WORKSHEET_COLUMNS:,}。",
                    )
                if not isinstance(value, dict):
                    raise DeliverableEditError(
                        "sheet_payload_invalid",
                        "单元格数据必须是对象。",
                    )
                populated += 1
        if populated > MAX_POPULATED_CELLS:
            raise DeliverableEditError(
                "too_many_cells",
                f"有效单元格数量超过上限 {MAX_POPULATED_CELLS:,}。",
            )
        for key, limit, label in (
            ("rowData", MAX_WORKSHEET_ROWS, "行"),
            ("columnData", MAX_WORKSHEET_COLUMNS, "列"),
        ):
            dimensions = sheet.get(key) or {}
            if not isinstance(dimensions, dict):
                raise DeliverableEditError(
                    "sheet_payload_invalid",
                    f"工作表包含无效的{label}属性。",
                )
            for raw_index in dimensions:
                try:
                    index = int(raw_index)
                except (TypeError, ValueError) as exc:
                    raise DeliverableEditError(
                        "sheet_payload_invalid",
                        f"工作表包含无效的{label}索引。",
                    ) from exc
                if index < 0 or index >= limit:
                    raise DeliverableEditError(
                        "sheet_bounds_invalid",
                        f"{label}索引超出 XLSX 上限。",
                    )
        merges = sheet.get("mergeData") or []
        if not isinstance(merges, list) or len(merges) > MAX_MERGED_RANGES:
            raise DeliverableEditError(
                "sheet_merge_invalid",
                f"合并区域数量超过上限 {MAX_MERGED_RANGES:,}，或数据结构无效。",
            )
        for merged in merges:
            if not isinstance(merged, dict):
                raise DeliverableEditError("sheet_merge_invalid", "合并区域数据无效。")
            try:
                start_row = int(merged.get("startRow"))
                start_column = int(merged.get("startColumn"))
                end_row = int(merged.get("endRow"))
                end_column = int(merged.get("endColumn"))
            except (TypeError, ValueError) as exc:
                raise DeliverableEditError(
                    "sheet_merge_invalid",
                    "合并区域坐标无效。",
                ) from exc
            if (
                start_row < 0
                or start_column < 0
                or end_row < start_row
                or end_column < start_column
                or end_row >= MAX_WORKSHEET_ROWS
                or end_column >= MAX_WORKSHEET_COLUMNS
            ):
                raise DeliverableEditError(
                    "sheet_merge_invalid",
                    "合并区域超出 XLSX 边界。",
                )
    return normalized_order, sheets


def univer_snapshot_to_xlsx(snapshot: dict[str, Any]) -> bytes:
    import openpyxl

    order, sheets = _validate_snapshot(snapshot)
    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)
    styles = snapshot.get("styles") if isinstance(snapshot.get("styles"), dict) else {}
    for sheet_id in order:
        sheet = sheets[sheet_id]
        title = str(sheet.get("name") or "Sheet")[:31] or "Sheet"
        worksheet = workbook.create_sheet(title=title)
        worksheet.sheet_state = "hidden" if int(sheet.get("hidden") or 0) else "visible"
        worksheet.sheet_view.showGridLines = bool(int(sheet.get("showGridlines", 1)))
        worksheet.sheet_view.rightToLeft = bool(int(sheet.get("rightToLeft", 0)))
        tab_color = str(sheet.get("tabColor") or "")
        if re.fullmatch(r"#[0-9A-Fa-f]{6}", tab_color):
            worksheet.sheet_properties.tabColor = "FF" + tab_color[1:].upper()
        freeze = sheet.get("freeze") if isinstance(sheet.get("freeze"), dict) else {}
        frozen_row = int(freeze.get("ySplit") or 0)
        frozen_column = int(freeze.get("xSplit") or 0)
        if frozen_row or frozen_column:
            worksheet.freeze_panes = worksheet.cell(
                row=frozen_row + 1,
                column=frozen_column + 1,
            )
        cell_data = sheet.get("cellData") or {}
        for row_key, row in cell_data.items():
            if not isinstance(row, dict):
                continue
            row_index = int(row_key) + 1
            for column_key, payload in row.items():
                if not isinstance(payload, dict):
                    continue
                column_index = int(column_key) + 1
                cell = worksheet.cell(row=row_index, column=column_index)
                formula = payload.get("f")
                if formula not in {None, ""}:
                    formula_text = str(formula)
                    cell.value = formula_text if formula_text.startswith("=") else "=" + formula_text
                else:
                    value = payload.get("v")
                    if int(payload.get("t") or 0) == 3:
                        value = bool(value)
                    cell.value = value
                style_ref = payload.get("s")
                style = styles.get(style_ref, {}) if isinstance(style_ref, str) else style_ref
                if isinstance(style, dict) and style:
                    _apply_univer_style(cell, style)
        for row_key, row_data in (sheet.get("rowData") or {}).items():
            if not isinstance(row_data, dict):
                continue
            dimension = worksheet.row_dimensions[int(row_key) + 1]
            if row_data.get("h") is not None:
                dimension.height = max(0, float(row_data["h"]) * 0.75)
            dimension.hidden = bool(int(row_data.get("hd") or 0))
        for column_key, column_data in (sheet.get("columnData") or {}).items():
            if not isinstance(column_data, dict):
                continue
            from openpyxl.utils.cell import get_column_letter

            dimension = worksheet.column_dimensions[get_column_letter(int(column_key) + 1)]
            if column_data.get("w") is not None:
                dimension.width = max(0, (float(column_data["w"]) - 5) / 7)
            dimension.hidden = bool(int(column_data.get("hd") or 0))
        for merged in sheet.get("mergeData") or []:
            if not isinstance(merged, dict):
                continue
            worksheet.merge_cells(
                start_row=int(merged.get("startRow") or 0) + 1,
                start_column=int(merged.get("startColumn") or 0) + 1,
                end_row=int(merged.get("endRow") or 0) + 1,
                end_column=int(merged.get("endColumn") or 0) + 1,
            )
    if workbook.worksheets and all(sheet.sheet_state != "visible" for sheet in workbook.worksheets):
        raise DeliverableEditError("all_sheets_hidden", "工作簿必须至少保留一个可见工作表。")
    calculation = getattr(workbook, "calculation", None)
    if calculation is not None:
        calculation.fullCalcOnLoad = True
        calculation.forceFullCalc = True
        calculation.calcMode = "auto"
    output = io.BytesIO()
    workbook.save(output)
    workbook.close()
    data = output.getvalue()
    try:
        check = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=False)
        check.close()
    except Exception as exc:
        raise DeliverableEditError("xlsx_export_invalid", f"导出的 XLSX 无法重新打开：{exc}") from exc
    return data


def univer_snapshot_to_rows(snapshot: dict[str, Any]) -> list[list[str]]:
    order, sheets = _validate_snapshot(snapshot)
    if len(order) != 1:
        raise DeliverableEditError(
            "tabular_multiple_sheets",
            "CSV/TSV 只能保存一个工作表；请删除新增工作表或另存为 XLSX。",
        )
    sheet = sheets[order[0]]
    cell_data = sheet.get("cellData") or {}
    max_row = -1
    max_column = -1
    for row_key, row in cell_data.items():
        if not isinstance(row, dict):
            continue
        max_row = max(max_row, int(row_key))
        if row:
            max_column = max(max_column, max(int(column) for column in row))
    if max_row < 0 or max_column < 0:
        return []
    rows = [["" for _column in range(max_column + 1)] for _row in range(max_row + 1)]
    for row_key, row in cell_data.items():
        if not isinstance(row, dict):
            continue
        for column_key, payload in row.items():
            if not isinstance(payload, dict):
                continue
            value = payload.get("f") if payload.get("f") not in {None, ""} else payload.get("v")
            if value is None:
                value = ""
            if isinstance(value, bool):
                value = "TRUE" if value else "FALSE"
            rows[int(row_key)][int(column_key)] = str(value)
    return rows


def load_editor_payload(session: EditSession) -> dict[str, Any]:
    extension = os.path.splitext(session.path)[1].lower()
    if session.descriptor.kind == "docx":
        return {"kind": "docx", "binary": Path(session.path).read_bytes()}
    if extension == ".xlsx":
        return {"kind": "sheet", "snapshot": xlsx_to_univer_snapshot(session.path)}
    raw = Path(session.path).read_bytes()
    bom = bytes.fromhex(session.bom_hex) if session.bom_hex else b""
    if bom and raw.startswith(bom):
        raw = raw[len(bom) :]
    text, _encoding, _bom, _newline = decode_text_bytes(raw, session.encoding)
    if extension in TABULAR_EXTENSIONS:
        delimiter = "\t" if extension == ".tsv" else ","
        rows = list(csv.reader(io.StringIO(text, newline=""), delimiter=delimiter))
        return {
            "kind": "sheet",
            "snapshot": rows_to_univer_snapshot(rows, os.path.basename(session.path)),
        }
    if session.descriptor.kind == "html":
        editable, preserved = prepare_html_edit_source(
            text,
            Path(session.path).parent.as_uri() + "/",
        )
        session.metadata["html_preserved_nodes"] = preserved
        session.metadata["html_original_source"] = text
        return {"kind": "html", "content": editable}
    return {"kind": "text", "content": text}


def _strict_json_loads(content: str) -> Any:
    def reject_constant(value):
        raise ValueError(f"JSON 不允许常量 {value}")

    def reject_duplicate_keys(pairs):
        result = {}
        for key, value in pairs:
            if key in result:
                raise ValueError(f"JSON 对象包含重复键 {key!r}")
            result[key] = value
        return result

    return json.loads(
        content,
        parse_constant=reject_constant,
        object_pairs_hook=reject_duplicate_keys,
    )


def validate_text_content(extension: str, content: str) -> None:
    extension = str(extension or "").lower()
    try:
        if extension == ".json":
            _strict_json_loads(content)
        elif extension in {".jsonl", ".ndjson"}:
            for line_number, line in enumerate(content.splitlines(), 1):
                if not line.strip():
                    continue
                try:
                    _strict_json_loads(line)
                except Exception as exc:
                    raise ValueError(f"第 {line_number} 行无效：{exc}") from exc
        elif extension == ".xml":
            if re.search(r"<!\s*(?:DOCTYPE|ENTITY)\b", content, flags=re.IGNORECASE):
                raise ValueError("XML 编辑文件不允许 DTD 或实体声明")
            ElementTree.fromstring(content)
        elif extension in {".yaml", ".yml"}:
            try:
                import yaml
            except ImportError as exc:
                raise DeliverableEditError(
                    "yaml_dependency_missing",
                    "缺少 PyYAML，无法验证 YAML 后保存。",
                ) from exc
            list(yaml.safe_load_all(content))
    except DeliverableEditError:
        raise
    except Exception as exc:
        label = extension.lstrip(".").upper()
        raise DeliverableEditError(
            "structured_text_invalid",
            f"{label} 语法验证失败：{exc}",
        ) from exc


def validate_docx_bytes(data: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            names = set(archive.namelist())
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise ValueError("缺少 DOCX 必需部件")
        from docx import Document

        Document(io.BytesIO(data))
    except Exception as exc:
        raise DeliverableEditError("docx_export_invalid", f"导出的 DOCX 无法重新打开：{exc}") from exc


def _next_docx_relationship_id(used_ids: set[str]) -> str:
    index = 1
    while True:
        candidate = f"rId{index}"
        if candidate not in used_ids:
            used_ids.add(candidate)
            return candidate
        index += 1


def _next_preserved_docx_part_name(kind: str, used_names: set[str]) -> str:
    index = 1
    while True:
        candidate = f"word/cowork-preserved-{kind}{index}.xml"
        if candidate not in used_names:
            used_names.add(candidate)
            return candidate
        index += 1


def restore_docx_header_footer(data: bytes, plan: dict[str, Any]) -> bytes:
    """Attach preserved, relationship-free header/footer parts to an exported body."""
    relationships = plan.get("relationships") if isinstance(plan, dict) else None
    preserved_parts = plan.get("parts") if isinstance(plan, dict) else None
    section_xml = plan.get("section_xml") if isinstance(plan, dict) else None
    if (
        not isinstance(relationships, dict)
        or not relationships
        or not isinstance(preserved_parts, dict)
        or not isinstance(section_xml, (bytes, bytearray))
    ):
        raise DeliverableEditError(
            "docx_preservation_invalid",
            "DOCX 页眉页脚保留信息无效，不能安全保存。",
        )

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as source:
            source_names = {
                name.replace("\\", "/") for name in source.namelist()
            }
            required_parts = {
                "[Content_Types].xml",
                "word/document.xml",
                "word/_rels/document.xml.rels",
            }
            missing = sorted(required_parts - source_names)
            if missing:
                raise DeliverableEditError(
                    "docx_export_invalid",
                    "导出的 DOCX 缺少必需部件：" + "、".join(missing),
                )

            document_dom = minidom.parseString(source.read("word/document.xml"))
            section_nodes = document_dom.getElementsByTagNameNS(
                _DOCX_WORDPROCESSING_NS,
                "sectPr",
            )
            if len(section_nodes) != 1:
                raise DeliverableEditError(
                    "docx_export_section_invalid",
                    "导出的 DOCX 页面结构异常，不能安全还原页眉页脚。",
                )

            preserved_section_dom = minidom.parseString(bytes(section_xml))
            preserved_section = preserved_section_dom.documentElement
            preserved_reference_nodes = [
                node
                for local_name in ("headerReference", "footerReference")
                for node in preserved_section.getElementsByTagNameNS(
                    _DOCX_WORDPROCESSING_NS,
                    local_name,
                )
            ]
            if not preserved_reference_nodes:
                raise DeliverableEditError(
                    "docx_preservation_invalid",
                    "DOCX 页眉页脚保留信息不完整。",
                )

            relationships_dom = minidom.parseString(
                source.read("word/_rels/document.xml.rels")
            )
            relationships_root = relationships_dom.documentElement
            existing_relationship_nodes = relationships_dom.getElementsByTagNameNS(
                _DOCX_PACKAGE_REL_NS,
                "Relationship",
            )
            used_relationship_ids = {
                node.getAttribute("Id")
                for node in existing_relationship_nodes
                if node.getAttribute("Id")
            }
            used_part_names = set(source_names)
            source_to_exported_part: dict[str, str] = {}
            added_parts: dict[str, bytes] = {}
            for source_part, content in sorted(preserved_parts.items()):
                source_part = str(source_part or "")
                if not isinstance(content, (bytes, bytearray)):
                    raise DeliverableEditError(
                        "docx_preservation_invalid",
                        "DOCX 页眉页脚部件内容无效。",
                    )
                kinds = {
                    str(item.get("kind") or "")
                    for item in relationships.values()
                    if isinstance(item, dict)
                    and str(item.get("part_name") or "") == source_part
                }
                if len(kinds) != 1 or next(iter(kinds)) not in {"header", "footer"}:
                    raise DeliverableEditError(
                        "docx_preservation_invalid",
                        "DOCX 页眉页脚部件类型无效。",
                    )
                kind = next(iter(kinds))
                exported_part = _next_preserved_docx_part_name(
                    kind,
                    used_part_names,
                )
                source_to_exported_part[source_part] = exported_part
                added_parts[exported_part] = bytes(content)

            relationship_id_map: dict[str, str] = {}
            for original_id, item in sorted(relationships.items()):
                if not isinstance(item, dict):
                    raise DeliverableEditError(
                        "docx_preservation_invalid",
                        "DOCX 页眉页脚关系信息无效。",
                    )
                kind = str(item.get("kind") or "")
                source_part = str(item.get("part_name") or "")
                exported_part = source_to_exported_part.get(source_part)
                if kind not in {"header", "footer"} or not exported_part:
                    raise DeliverableEditError(
                        "docx_preservation_invalid",
                        "DOCX 页眉页脚关系目标无效。",
                    )
                relationship_id = _next_docx_relationship_id(
                    used_relationship_ids
                )
                relationship_id_map[str(original_id)] = relationship_id
                relationship_node = relationships_dom.createElementNS(
                    _DOCX_PACKAGE_REL_NS,
                    "Relationship",
                )
                relationship_node.setAttribute("Id", relationship_id)
                relationship_node.setAttribute(
                    "Type",
                    _DOCX_HEADER_REL_TYPE
                    if kind == "header"
                    else _DOCX_FOOTER_REL_TYPE,
                )
                relationship_node.setAttribute(
                    "Target",
                    posixpath.relpath(exported_part, "word"),
                )
                relationships_root.appendChild(relationship_node)

            for reference in preserved_reference_nodes:
                original_id = reference.getAttributeNS(
                    _DOCX_OFFICE_REL_NS,
                    "id",
                )
                relationship_id = relationship_id_map.get(original_id)
                if not relationship_id:
                    raise DeliverableEditError(
                        "docx_preservation_invalid",
                        "DOCX 页眉页脚引用无法还原。",
                    )
                reference.setAttributeNS(
                    _DOCX_OFFICE_REL_NS,
                    "r:id",
                    relationship_id,
                )

            exported_section = section_nodes[0]
            exported_section.parentNode.replaceChild(
                document_dom.importNode(preserved_section, deep=True),
                exported_section,
            )

            content_types_dom = minidom.parseString(
                source.read("[Content_Types].xml")
            )
            content_types_root = content_types_dom.documentElement
            for part_name in sorted(added_parts):
                kind = "header" if "-header" in part_name else "footer"
                override = content_types_dom.createElementNS(
                    _DOCX_CONTENT_TYPES_NS,
                    "Override",
                )
                override.setAttribute("PartName", f"/{part_name}")
                override.setAttribute(
                    "ContentType",
                    _DOCX_HEADER_CONTENT_TYPE
                    if kind == "header"
                    else _DOCX_FOOTER_CONTENT_TYPE,
                )
                content_types_root.appendChild(override)

            replacements = {
                "[Content_Types].xml": content_types_dom.toxml(
                    encoding="utf-8"
                ),
                "word/document.xml": document_dom.toxml(encoding="utf-8"),
                "word/_rels/document.xml.rels": relationships_dom.toxml(
                    encoding="utf-8"
                ),
            }
            output = io.BytesIO()
            with zipfile.ZipFile(
                output,
                "w",
                compression=zipfile.ZIP_DEFLATED,
            ) as target:
                target.comment = source.comment
                for item in source.infolist():
                    normalized_name = item.filename.replace("\\", "/")
                    target.writestr(
                        item,
                        replacements.get(normalized_name, source.read(item.filename)),
                    )
                for part_name, content in sorted(added_parts.items()):
                    target.writestr(
                        part_name,
                        content,
                        compress_type=zipfile.ZIP_DEFLATED,
                    )
    except DeliverableEditError:
        raise
    except (zipfile.BadZipFile, ExpatError, ValueError, KeyError) as exc:
        raise DeliverableEditError(
            "docx_preservation_failed",
            f"无法还原 DOCX 页眉页脚：{exc}",
        ) from exc
    return output.getvalue()


def serialize_editor_payload(session: EditSession, payload: Any) -> bytes:
    extension = os.path.splitext(session.path)[1].lower()
    if session.descriptor.kind == "docx":
        if not isinstance(payload, (bytes, bytearray)):
            raise DeliverableEditError("docx_payload_invalid", "DOCX 编辑器没有返回有效文档。")
        data = bytes(payload)
        header_footer_plan = session.metadata.get("_docx_header_footer_plan")
        if header_footer_plan is not None:
            data = restore_docx_header_footer(data, header_footer_plan)
        if len(data) > OFFICE_MAX_BYTES:
            raise DeliverableEditError("file_too_large", "编辑后的 DOCX 超过 25 MiB。")
        validate_docx_bytes(data)
        return data
    if extension == ".xlsx":
        if not isinstance(payload, dict):
            raise DeliverableEditError("sheet_payload_invalid", "表格编辑器没有返回有效工作簿。")
        data = univer_snapshot_to_xlsx(payload)
        if len(data) > OFFICE_MAX_BYTES:
            raise DeliverableEditError("file_too_large", "编辑后的 XLSX 超过 25 MiB。")
        return data
    if extension in TABULAR_EXTENSIONS:
        if not isinstance(payload, dict):
            raise DeliverableEditError("sheet_payload_invalid", "表格编辑器没有返回有效数据。")
        rows = univer_snapshot_to_rows(payload)
        stream = io.StringIO(newline="")
        csv.writer(
            stream,
            delimiter="\t" if extension == ".tsv" else ",",
            lineterminator=session.newline,
        ).writerows(rows)
        data = encode_text_content(stream.getvalue(), session)
        if len(data) > TEXT_MAX_BYTES:
            raise DeliverableEditError("file_too_large", "编辑后的表格文本超过 10 MiB。")
        return data
    if not isinstance(payload, str):
        raise DeliverableEditError("text_payload_invalid", "文本编辑器没有返回有效内容。")
    if session.descriptor.kind == "html":
        payload = apply_html_dom_patch(
            str(session.metadata.get("html_original_source") or ""),
            payload,
            list(session.metadata.get("html_preserved_nodes") or []),
        )
    validate_text_content(extension, payload)
    data = encode_text_content(payload, session)
    if len(data) > TEXT_MAX_BYTES:
        raise DeliverableEditError("file_too_large", "编辑后的文本超过 10 MiB。")
    return data


def backup_paths(path: str, backup_root: str | None = None) -> tuple[str, str]:
    normalized = os.path.normcase(os.path.abspath(path))
    key = hashlib.sha256(normalized.encode("utf-8")).hexdigest()
    root = os.path.abspath(
        backup_root or os.path.join(get_app_data_dir(), "deliverable_backups")
    )
    directory = os.path.join(root, key)
    extension = os.path.splitext(path)[1].lower()
    return os.path.join(directory, f"previous{extension}"), os.path.join(directory, "metadata.json")


def _write_bytes_to_temp(directory: str, prefix: str, data: bytes) -> str:
    descriptor, temp_path = tempfile.mkstemp(prefix=prefix, suffix=".tmp", dir=directory)
    try:
        with os.fdopen(descriptor, "wb") as handle:
            handle.write(data)
            handle.flush()
            os.fsync(handle.fileno())
    except Exception:
        try:
            os.unlink(temp_path)
        except OSError:
            pass
        raise
    return temp_path


def _write_json_atomic(path: str, payload: dict[str, Any]) -> None:
    directory = os.path.dirname(path)
    os.makedirs(directory, exist_ok=True)
    data = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
    temp_path = _write_bytes_to_temp(directory, ".metadata-", data)
    try:
        os.replace(temp_path, path)
    finally:
        if os.path.exists(temp_path):
            os.unlink(temp_path)


def _validate_serialized_path(path: str, extension: str, selected_encoding: str = "") -> None:
    data = Path(path).read_bytes()
    if extension == ".docx":
        report = _preflight_docx(path, EDITOR_BY_EXTENSION[".docx"])
        if not report.allowed:
            raise DeliverableEditError(
                "docx_export_invalid",
                "临时 DOCX 未通过兼容性验证：\n" + report.message,
            )
    elif extension == ".xlsx":
        report = _preflight_xlsx(path, EDITOR_BY_EXTENSION[".xlsx"])
        if not report.allowed:
            raise DeliverableEditError(
                "xlsx_export_invalid",
                "临时 XLSX 未通过兼容性验证：\n" + report.message,
            )
    elif extension in EDITABLE_EXTENSIONS:
        if not data:
            return
        text, _encoding, _bom, _newline = decode_text_bytes(data, selected_encoding)
        validate_text_content(extension, text)
        if extension in TABULAR_EXTENSIONS:
            delimiter = "\t" if extension == ".tsv" else ","
            try:
                rows = csv.reader(io.StringIO(text, newline=""), delimiter=delimiter)
                populated = sum(1 for row in rows for cell in row if cell != "")
            except csv.Error as exc:
                raise DeliverableEditError(
                    "tabular_parse_failed",
                    f"临时表格文本无法解析：{exc}",
                ) from exc
            if populated > MAX_POPULATED_CELLS:
                raise DeliverableEditError(
                    "too_many_cells",
                    f"有效单元格数量超过上限 {MAX_POPULATED_CELLS:,}。",
                )


def atomic_save_session(
    session: EditSession,
    data: bytes,
    backup_root: str | None = None,
    event_logger: Callable[..., None] | None = None,
) -> SaveResult:
    path = os.path.abspath(session.path)
    logger = event_logger or (lambda *_args, **_kwargs: None)
    if not os.path.isfile(path):
        raise DeliverableEditError("file_not_found", "原文件已被删除，不能保存。")
    current_fingerprint = sha256_file(path)
    if current_fingerprint != session.initial_fingerprint:
        raise ExternalModificationError()
    if len(data) > session.descriptor.max_bytes:
        raise DeliverableEditError(
            "file_too_large",
            f"编辑后的文件超过 {session.descriptor.max_bytes // MIB} MiB 保存上限。",
        )
    parent = os.path.dirname(path)
    extension = os.path.splitext(path)[1].lower()
    temp_path = ""
    backup_temp = ""
    backup_path, metadata_path = backup_paths(path, backup_root)
    try:
        logger("deliverable_edit_save_start", path=path, bytes=len(data))
        temp_path = _write_bytes_to_temp(parent, f".{os.path.basename(path)}.cowork-edit-", data)
        _validate_serialized_path(temp_path, extension, session.encoding)
        backup_dir = os.path.dirname(backup_path)
        os.makedirs(backup_dir, exist_ok=True)
        backup_temp = os.path.join(backup_dir, f".previous-{uuid.uuid4().hex}.tmp")
        shutil.copy2(path, backup_temp)
        with open(backup_temp, "r+b") as handle:
            os.fsync(handle.fileno())
        if sha256_file(backup_temp) != current_fingerprint:
            raise ExternalModificationError("文件在创建备份时发生变化，不能覆盖。")
        if sha256_file(path) != current_fingerprint:
            raise ExternalModificationError("文件在保存过程中被其他程序修改，不能覆盖。")
        os.replace(backup_temp, backup_path)
        backup_temp = ""
        _write_json_atomic(
            metadata_path,
            {
                "source_path": path,
                "source_key": os.path.normcase(path),
                "source_fingerprint": current_fingerprint,
                "saved_at": datetime.now().astimezone().isoformat(),
                "backup_path": backup_path,
                "encoding": session.encoding,
                "newline": session.newline,
                "bom_hex": session.bom_hex,
            },
        )
        logger("deliverable_edit_backup_finish", path=path, backup_path=backup_path)
        if sha256_file(path) != current_fingerprint:
            raise ExternalModificationError("文件在保存过程中被其他程序修改，不能覆盖。")
        fingerprint = hashlib.sha256(data).hexdigest()
        os.replace(temp_path, path)
        temp_path = ""
        session.initial_fingerprint = fingerprint
        session.initial_size = len(data)
        session.dirty = False
        session.state = "saved"
        logger("deliverable_edit_save_finish", path=path, bytes=len(data))
        return SaveResult(
            path=path,
            backup_path=backup_path,
            fingerprint=fingerprint,
            bytes_written=len(data),
        )
    except DeliverableEditError:
        logger("deliverable_edit_save_error", path=path)
        raise
    except Exception as exc:
        logger("deliverable_edit_save_error", path=path, error=str(exc))
        raise DeliverableEditError("atomic_save_failed", f"保存文件失败：{exc}") from exc
    finally:
        for candidate in (temp_path, backup_temp):
            if candidate and os.path.exists(candidate):
                try:
                    os.unlink(candidate)
                except OSError:
                    pass


def save_copy(
    session: EditSession,
    data: bytes,
    target_path: str,
) -> SaveResult:
    target = os.path.abspath(os.path.normpath(str(target_path or "")))
    if not target:
        raise DeliverableEditError("copy_path_missing", "没有选择另存位置。")
    if os.path.normcase(target) == os.path.normcase(session.path):
        raise DeliverableEditError("copy_matches_source", "另存位置不能与原文件相同。")
    if os.path.exists(target):
        raise DeliverableEditError("copy_exists", "目标文件已存在，请选择新的文件名。")
    if os.path.splitext(target)[1].lower() != os.path.splitext(session.path)[1].lower():
        raise DeliverableEditError("copy_extension_mismatch", "另存文件必须保持原扩展名。")
    if len(data) > session.descriptor.max_bytes:
        raise DeliverableEditError(
            "file_too_large",
            f"另存文件超过 {session.descriptor.max_bytes // MIB} MiB 保存上限。",
        )
    os.makedirs(os.path.dirname(target), exist_ok=True)
    extension = os.path.splitext(target)[1].lower()
    temp_path = _write_bytes_to_temp(
        os.path.dirname(target),
        f".{os.path.basename(target)}.cowork-copy-",
        data,
    )
    try:
        _validate_serialized_path(temp_path, extension, session.encoding)
        os.replace(temp_path, target)
        temp_path = ""
        return SaveResult(
            path=target,
            backup_path="",
            fingerprint=sha256_file(target),
            bytes_written=len(data),
        )
    except DeliverableEditError:
        raise
    except Exception as exc:
        raise DeliverableEditError("save_copy_failed", f"另存文件失败：{exc}") from exc
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except OSError:
                pass


def restore_previous_version(
    path: str,
    backup_root: str | None = None,
) -> SaveResult:
    source, _stat = _require_regular_file(path)
    previous_path, metadata_path = backup_paths(source, backup_root)
    if not os.path.isfile(previous_path) or not os.path.isfile(metadata_path):
        raise DeliverableEditError("backup_not_found", "当前文件没有可恢复的上一版。")
    try:
        metadata = json.loads(Path(metadata_path).read_text(encoding="utf-8"))
    except Exception as exc:
        raise DeliverableEditError("backup_metadata_invalid", f"备份元数据损坏：{exc}") from exc
    if os.path.normcase(str(metadata.get("source_path") or "")) != os.path.normcase(source):
        raise DeliverableEditError("backup_mismatch", "备份与当前文件不匹配。")

    source_dir = os.path.dirname(source)
    backup_dir = os.path.dirname(previous_path)
    extension = os.path.splitext(source)[1].lower()
    restore_temp = ""
    current_temp = ""
    rollback_temp = ""
    try:
        restore_temp = os.path.join(
            source_dir,
            f".{os.path.basename(source)}.cowork-restore-{uuid.uuid4().hex}.tmp",
        )
        shutil.copy2(previous_path, restore_temp)
        if extension in {".docx", ".xlsx"}:
            _validate_serialized_path(restore_temp, extension)
        current_temp = os.path.join(backup_dir, f".current-{uuid.uuid4().hex}.tmp")
        shutil.copy2(source, current_temp)
        os.replace(restore_temp, source)
        restore_temp = ""
        try:
            os.replace(current_temp, previous_path)
            current_temp = ""
        except Exception as swap_exc:
            rollback_temp = os.path.join(
                source_dir,
                f".{os.path.basename(source)}.cowork-rollback-{uuid.uuid4().hex}.tmp",
            )
            shutil.copy2(current_temp, rollback_temp)
            os.replace(rollback_temp, source)
            rollback_temp = ""
            raise DeliverableEditError(
                "backup_swap_failed",
                f"恢复时无法更新上一版，已回滚原文件：{swap_exc}",
            ) from swap_exc
        fingerprint = sha256_file(source)
        _write_json_atomic(
            metadata_path,
            {
                "source_path": source,
                "source_key": os.path.normcase(source),
                "source_fingerprint": fingerprint,
                "saved_at": datetime.now().astimezone().isoformat(),
                "backup_path": previous_path,
            },
        )
        return SaveResult(
            path=source,
            backup_path=previous_path,
            fingerprint=fingerprint,
            bytes_written=os.path.getsize(source),
        )
    finally:
        for candidate in (restore_temp, current_temp, rollback_temp):
            if candidate and os.path.exists(candidate):
                try:
                    os.unlink(candidate)
                except OSError:
                    pass
