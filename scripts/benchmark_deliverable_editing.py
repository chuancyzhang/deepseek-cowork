"""Run the reference deliverable-editing size and latency checks.

The benchmark creates disposable fixtures: a 100-page, roughly 10 MiB DOCX
with one image per page and an XLSX with 100,000 populated cells.  It measures
Python preflight plus payload preparation and can optionally run the real
offline WebEngine round-trip against the large DOCX.
"""

from __future__ import annotations

import argparse
import io
import os
from pathlib import Path
import random
import subprocess
import sys
import tempfile
import time


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from core.deliverable_editing import create_edit_session, load_editor_payload


def build_docx(path: Path, pages: int = 100) -> None:
    from docx import Document
    from docx.shared import Inches
    from PIL import Image

    document = Document()
    for page in range(pages):
        document.add_heading(f"性能样本文档 · 第 {page + 1} 页", level=1)
        document.add_paragraph(
            "这是用于验证内嵌编辑性能预算的普通图文页面，包含标题、正文和一张图片。"
        )
        pixels = random.Random(page).randbytes(192 * 192 * 3)
        image = Image.frombytes("RGB", (192, 192), pixels)
        image_stream = io.BytesIO()
        image.save(image_stream, format="PNG", optimize=False)
        image_stream.seek(0)
        document.add_picture(image_stream, width=Inches(2.2))
        if page + 1 < pages:
            document.add_page_break()
    document.save(path)


def build_xlsx(path: Path, populated_cells: int = 100_000) -> None:
    import openpyxl

    columns = 10
    rows = populated_cells // columns
    workbook = openpyxl.Workbook(write_only=True)
    worksheet = workbook.create_sheet("十万单元格")
    for row in range(rows):
        worksheet.append([row * columns + column for column in range(columns)])
    workbook.save(path)


def measure(path: Path) -> tuple[float, float, dict]:
    started = time.perf_counter()
    session, report = create_edit_session(str(path))
    preflight_seconds = time.perf_counter() - started
    started = time.perf_counter()
    payload = load_editor_payload(session)
    payload_seconds = time.perf_counter() - started
    return preflight_seconds, payload_seconds, {
        "kind": payload.get("kind"),
        **report.metadata,
    }


def main(argv=None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--budget-seconds", type=float, default=8.0)
    parser.add_argument(
        "--with-webengine",
        action="store_true",
        help="Also load and export the large DOCX in the real offline editor.",
    )
    args = parser.parse_args(argv)

    with tempfile.TemporaryDirectory(prefix="cowork-editor-benchmark-") as directory:
        root = Path(directory)
        docx_path = root / "100-pages.docx"
        xlsx_path = root / "100k-cells.xlsx"
        build_docx(docx_path)
        build_xlsx(xlsx_path)

        failures = []
        for path in (docx_path, xlsx_path):
            preflight, payload, metadata = measure(path)
            total = preflight + payload
            print(
                f"{path.name}: size={path.stat().st_size / 1024 / 1024:.2f} MiB "
                f"preflight={preflight:.3f}s payload={payload:.3f}s total={total:.3f}s "
                f"metadata={metadata}"
            )
            if total > args.budget_seconds:
                failures.append(
                    f"{path.name} core load {total:.3f}s exceeds {args.budget_seconds:.3f}s"
                )

        if args.with_webengine:
            environment = os.environ.copy()
            environment.setdefault("QT_QPA_PLATFORM", "offscreen")
            environment.setdefault("QT_QUICK_BACKEND", "software")
            environment.setdefault("QTWEBENGINE_CHROMIUM_FLAGS", "--disable-gpu")
            completed = subprocess.run(
                [
                    sys.executable,
                    str(ROOT / "scripts" / "smoke_deliverable_editors.py"),
                    "--docx-path",
                    str(docx_path),
                    "--model-budget-seconds",
                    str(args.budget_seconds),
                ],
                cwd=ROOT,
                env=environment,
                capture_output=True,
                text=True,
                timeout=180,
            )
            if completed.stdout:
                print(completed.stdout.rstrip())
            if completed.stderr:
                print(completed.stderr.rstrip(), file=sys.stderr)
            if completed.returncode:
                failures.append(
                    f"WebEngine round-trip failed with exit code {completed.returncode}"
                )

        if failures:
            for failure in failures:
                print(f"[ERROR] {failure}", file=sys.stderr)
            return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
