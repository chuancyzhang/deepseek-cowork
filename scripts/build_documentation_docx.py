"""Generate the two user-facing DOCX guides from canonical Markdown.

The builder implements the documents skill's compact_reference_guide preset:
US Letter, 1 inch margins, Calibri 11 pt body, 1.25 line spacing, explicit
heading/list rhythm, and fixed-DXA table geometry. Chinese glyphs use the named
`cjk_font` override (Microsoft YaHei) while Latin text remains Calibri.
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import unquote

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image
except ImportError:  # pragma: no cover - dependency contract requires Pillow
    Image = None


ROOT = Path(__file__).resolve().parents[1]
APP_VERSION = "5.1.2"
PRESET_NAME = "compact_reference_guide"

BODY_ASCII_FONT = "Calibri"
BODY_CJK_FONT = "Microsoft YaHei"
MONO_FONT = "Consolas"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x20, 0x37, 0x54)
MUTED = RGBColor(0x64, 0x6B, 0x75)
BODY = RGBColor(0x20, 0x23, 0x29)
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
CODE_FILL = "F5F6F8"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

IMAGE_RE = re.compile(r"^!\[([^\]]*)]\(([^)]+)\)\s*$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
BULLET_RE = re.compile(r"^\s*[-*+]\s+(.+)$")
NUMBER_RE = re.compile(r"^\s*\d+[.)]\s+(.+)$")
LINK_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+]\([^)]+\))")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*$")


@dataclass(frozen=True)
class DocumentSpec:
    source: Path
    output: Path
    title: str
    subtitle: str
    audience: str
    running_label: str


SPECS = {
    "user-guide": DocumentSpec(
        source=ROOT / "docs" / "user-guide.md",
        output=ROOT / "USER_GUIDE.docx",
        title="DeepSeek Cowork 用户指南",
        subtitle="从第一次打开，到让 AI 帮你完成真实工作",
        audience="面向第一次使用桌面 Agent 的普通用户",
        running_label="DeepSeek Cowork 用户指南",
    ),
    "theme-visualize": DocumentSpec(
        source=ROOT / "docs" / "guides" / "ai-theme-and-visualize.md",
        output=ROOT / "AI主题与Visualize普通用户指南.docx",
        title="让软件更像你，也让答案动起来",
        subtitle="从 AI 自定义主题到 Visualize 交互可视化",
        audience="面向希望个性化界面和探索交互答案的普通用户",
        running_label="AI 主题与 Visualize",
    ),
}


def _set_run_font(run, *, ascii_font=BODY_ASCII_FONT, cjk_font=BODY_CJK_FONT):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), cjk_font)


def _set_cell_margins(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa: list[int]):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def _set_paragraph_fill(paragraph, color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), color)


def _add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    _set_run_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def _configure_styles(document: Document):
    normal = document.styles["Normal"]
    normal.font.name = BODY_ASCII_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_ASCII_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_ASCII_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CJK_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = BODY_ASCII_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = BODY_ASCII_FONT
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CJK_FONT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in document.styles:
        style = document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = document.styles["Code Block"]
    style.font.name = MONO_FONT
    style.font.size = Pt(9)
    style._element.rPr.rFonts.set(qn("w:ascii"), MONO_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), MONO_FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CJK_FONT)
    style.paragraph_format.left_indent = Inches(0.14)
    style.paragraph_format.right_indent = Inches(0.14)
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.0

    if "Figure Caption" not in document.styles:
        style = document.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = document.styles["Figure Caption"]
    style.font.name = BODY_ASCII_FONT
    style.font.size = Pt(9)
    style.font.italic = True
    style.font.color.rgb = MUTED
    style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CJK_FONT)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_before = Pt(3)
    style.paragraph_format.space_after = Pt(10)
    style.paragraph_format.keep_with_next = False


def _configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def _configure_header_footer(section, label):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(label)
    _set_run_font(run)
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    run = paragraph.add_run(f"DeepSeek Cowork · {APP_VERSION}")
    _set_run_font(run)
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED
    paragraph.add_run("\t")
    page_label = paragraph.add_run("第 ")
    _set_run_font(page_label)
    page_label.font.size = Pt(9)
    page_label.font.color.rgb = MUTED
    _add_page_field(paragraph)
    suffix = paragraph.add_run(" 页")
    _set_run_font(suffix)
    suffix.font.size = Pt(9)
    suffix.font.color.rgb = MUTED


def _add_cover(document: Document, spec: DocumentSpec):
    # editorial_cover pattern with a compact vertical budget for image-heavy guides.
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(44)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("DEEPSEEK COWORK · REFERENCE GUIDE")
    _set_run_font(run)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run(spec.title)
    _set_run_font(run)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run(spec.subtitle)
    _set_run_font(run)
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(6)
    run = meta.add_run(f"适用版本 {APP_VERSION}")
    _set_run_font(run)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BODY

    audience = document.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = audience.add_run(spec.audience)
    _set_run_font(run)
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    document.add_page_break()


def _add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), BODY_ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), BODY_CJK_FONT)
    rpr.extend([rfonts, color, underline])
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _append_inline(paragraph, text: str):
    position = 0
    for match in LINK_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            _set_run_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run)
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run)
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, ascii_font=MONO_FONT)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_BLUE
        else:
            link_match = re.match(r"\[([^\]]+)]\(([^)]+)\)", token)
            label, target = link_match.groups()
            if target.startswith(("http://", "https://", "mailto:")):
                _add_hyperlink(paragraph, label, target)
            else:
                run = paragraph.add_run(label)
                _set_run_font(run)
                run.font.color.rgb = BLUE
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        _set_run_font(run)


def _add_paragraph(document, text, *, style=None, fill=None):
    paragraph = document.add_paragraph(style=style)
    _append_inline(paragraph, text.strip())
    if fill:
        _set_paragraph_fill(paragraph, fill)
        paragraph.paragraph_format.left_indent = Inches(0.14)
        paragraph.paragraph_format.right_indent = Inches(0.14)
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(8)
    return paragraph


def _parse_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _column_widths(rows: list[list[str]]) -> list[int]:
    column_count = max(len(row) for row in rows)
    weights = []
    for index in range(column_count):
        longest = max((len(row[index]) if index < len(row) else 0) for row in rows)
        weights.append(max(8, min(longest, 48)))
    total_weight = sum(weights)
    widths = [max(1200, round(CONTENT_WIDTH_DXA * weight / total_weight)) for weight in weights]
    difference = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += difference
    if widths[-1] < 1200:
        deficit = 1200 - widths[-1]
        donor = max(range(len(widths) - 1), key=lambda idx: widths[idx])
        widths[donor] -= deficit
        widths[-1] = 1200
    return widths


def _add_table(document, rows: list[list[str]]):
    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=column_count)
    widths = _column_widths(normalized)
    _set_table_geometry(table, widths)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for row_index, row in enumerate(normalized):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            _append_inline(paragraph, value)
            if row_index == 0:
                _set_cell_fill(cell, TABLE_FILL)
                for run in paragraph.runs:
                    run.bold = True
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _add_image(document, source: Path, alt: str, target: str):
    image_path = (source.parent / unquote(target.strip().strip("<>"))).resolve()
    if not image_path.exists():
        raise FileNotFoundError(f"missing image referenced by {source}: {target}")
    width_inches = 6.5
    height_inches = 7.45
    if Image is not None:
        with Image.open(image_path) as image:
            pixel_width, pixel_height = image.size
        ratio = pixel_width / max(pixel_height, 1)
        if width_inches / ratio > height_inches:
            width_inches = height_inches * ratio
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    if alt:
        caption = document.add_paragraph(alt, style="Figure Caption")
        caption.paragraph_format.keep_together = True


def _is_special(line: str, next_line: str = "") -> bool:
    stripped = line.strip()
    return bool(
        not stripped
        or stripped.startswith("```")
        or HEADING_RE.match(stripped)
        or IMAGE_RE.match(stripped)
        or BULLET_RE.match(stripped)
        or NUMBER_RE.match(stripped)
        or stripped.startswith(">")
        or (stripped.startswith("|") and TABLE_SEPARATOR_RE.match(next_line.strip()))
    )


def _render_markdown(document: Document, source: Path):
    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    first_h1_skipped = False
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            index += 1
            content = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                content.append(lines[index])
                index += 1
            index += 1
            if language:
                label = document.add_paragraph()
                label.paragraph_format.space_after = Pt(2)
                run = label.add_run(language.upper())
                _set_run_font(run, ascii_font=MONO_FONT)
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = MUTED
            paragraph = document.add_paragraph("\n".join(content), style="Code Block")
            _set_paragraph_fill(paragraph, CODE_FILL)
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level == 1 and not first_h1_skipped:
                first_h1_skipped = True
                index += 1
                continue
            target_level = min(3, max(1, level - 1 if level > 1 else 1))
            paragraph = document.add_paragraph(style=f"Heading {target_level}")
            _append_inline(paragraph, text)
            index += 1
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            _add_image(document, source, image_match.group(1), image_match.group(2))
            index += 1
            while index < len(lines) and not lines[index].strip():
                index += 1
            if index < len(lines) and re.match(r"^\*截图[^*]*\*$", lines[index].strip()):
                index += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip()[1:].strip())
                index += 1
            text = " ".join(part for part in quote_lines if part)
            if text:
                paragraph = _add_paragraph(document, text, fill=CALLOUT_FILL)
                paragraph.paragraph_format.keep_together = True
            continue

        next_line = lines[index + 1] if index + 1 < len(lines) else ""
        if stripped.startswith("|") and TABLE_SEPARATOR_RE.match(next_line.strip()):
            rows = [_parse_table_row(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(_parse_table_row(lines[index]))
                index += 1
            _add_table(document, rows)
            continue

        bullet_match = BULLET_RE.match(stripped)
        if bullet_match:
            _add_paragraph(document, bullet_match.group(1), style="List Bullet")
            index += 1
            continue

        number_match = NUMBER_RE.match(stripped)
        if number_match:
            _add_paragraph(document, number_match.group(1), style="List Number")
            index += 1
            continue

        paragraph_lines = [stripped.rstrip("\\").rstrip()]
        index += 1
        while index < len(lines):
            candidate = lines[index]
            following = lines[index + 1] if index + 1 < len(lines) else ""
            if _is_special(candidate, following):
                break
            paragraph_lines.append(candidate.strip().rstrip("\\").rstrip())
            index += 1
        _add_paragraph(document, " ".join(paragraph_lines))


def _set_document_metadata(document: Document, spec: DocumentSpec):
    props = document.core_properties
    props.title = spec.title
    props.subject = f"DeepSeek Cowork {APP_VERSION} 用户文档"
    props.author = "DeepSeek Cowork"
    props.keywords = f"DeepSeek Cowork, {PRESET_NAME}, {APP_VERSION}"
    props.comments = (
        f"Generated from {spec.source.relative_to(ROOT).as_posix()}; "
        f"preset={PRESET_NAME}; named_override=cjk_font:{BODY_CJK_FONT}"
    )


def build(spec: DocumentSpec):
    document = Document()
    _configure_styles(document)
    for section in document.sections:
        _configure_section(section)
        _configure_header_footer(section, spec.running_label)
    _set_document_metadata(document, spec)
    _add_cover(document, spec)
    _render_markdown(document, spec.source)
    spec.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(spec.output)
    audit(spec)
    return spec.output


def _integer_attr(element, name):
    if element is None:
        return None
    value = element.get(qn(name))
    return int(value) if value is not None else None


def audit(spec: DocumentSpec):
    """Fail generation when preset geometry or source coverage drifts."""
    document = Document(spec.output)
    errors = []
    for index, section in enumerate(document.sections, 1):
        expected = {
            "page_width": Inches(8.5),
            "page_height": Inches(11),
            "top_margin": Inches(1),
            "right_margin": Inches(1),
            "bottom_margin": Inches(1),
            "left_margin": Inches(1),
            "header_distance": Inches(0.492),
            "footer_distance": Inches(0.492),
        }
        for attribute, expected_value in expected.items():
            actual_value = getattr(section, attribute)
            # OOXML section distances are stored in whole twips; accept one-twip quantization.
            if abs(int(actual_value) - int(expected_value)) > 635:
                errors.append(
                    f"section {index} {attribute}: expected {int(expected_value)}, got {int(actual_value)}"
                )

    expected_styles = {
        "Normal": (11, 6),
        "Heading 1": (16, 10),
        "Heading 2": (13, 7),
        "Heading 3": (12, 5),
        "List Bullet": (11, 4),
        "List Number": (11, 4),
    }
    for style_name, (font_size, after) in expected_styles.items():
        style = document.styles[style_name]
        actual_size = style.font.size.pt if style.font.size else None
        actual_after = (
            style.paragraph_format.space_after.pt
            if style.paragraph_format.space_after is not None
            else None
        )
        if actual_size != font_size:
            errors.append(f"{style_name}: expected {font_size} pt, got {actual_size}")
        if actual_after != after:
            errors.append(f"{style_name}: expected {after} pt after, got {actual_after}")

    for table_index, table in enumerate(document.tables, 1):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
        if _integer_attr(tbl_w, "w:w") != CONTENT_WIDTH_DXA:
            errors.append(f"table {table_index}: tblW is not {CONTENT_WIDTH_DXA}")
        if _integer_attr(tbl_ind, "w:w") != TABLE_INDENT_DXA:
            errors.append(f"table {table_index}: tblInd is not {TABLE_INDENT_DXA}")
        grid_widths = [_integer_attr(node, "w:w") for node in table._tbl.tblGrid]
        if sum(width or 0 for width in grid_widths) != CONTENT_WIDTH_DXA:
            errors.append(f"table {table_index}: tblGrid does not total {CONTENT_WIDTH_DXA}")
        for row_index, row in enumerate(table.rows, 1):
            cell_widths = []
            for cell in row.cells:
                tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
                cell_widths.append(_integer_attr(tc_w, "w:w"))
            if cell_widths != grid_widths:
                errors.append(
                    f"table {table_index} row {row_index}: tcW {cell_widths} != grid {grid_widths}"
                )

    source_image_count = len(
        re.findall(r"^!\[[^\]]*]\([^)]+\)\s*$", spec.source.read_text(encoding="utf-8"), re.MULTILINE)
    )
    document_image_count = len(document.part._package.image_parts)
    if document_image_count != source_image_count:
        errors.append(
            f"image coverage: expected {source_image_count}, embedded {document_image_count}"
        )

    if PRESET_NAME not in (document.core_properties.keywords or ""):
        errors.append("preset metadata is missing")
    if errors:
        raise RuntimeError(f"DOCX audit failed for {spec.output.name}:\n- " + "\n- ".join(errors))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "targets",
        nargs="*",
        choices=tuple(SPECS) + ("all",),
        default=["all"],
    )
    args = parser.parse_args()
    target_names = list(SPECS) if "all" in args.targets else args.targets
    for target_name in target_names:
        output = build(SPECS[target_name])
        print(output)
    return 0


if __name__ == "__main__":
    sys.exit(main())
