"""Load and round-trip every offline deliverable editor in an isolated process."""

from __future__ import annotations

import argparse
import base64
import io
import json
import os
from pathlib import Path
import sys
import time

from PySide6.QtCore import QObject, QTimer, QUrl, Signal, Slot
from PySide6.QtWidgets import QApplication
from PySide6.QtWebChannel import QWebChannel
from PySide6.QtWebEngineCore import (
    QWebEnginePage,
    QWebEngineProfile,
    QWebEngineUrlRequestInterceptor,
)
from PySide6.QtWebEngineWidgets import QWebEngineView


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from core.deliverable_editing import (  # noqa: E402
    create_edit_session,
    rows_to_univer_snapshot,
    serialize_editor_payload,
    validate_docx_bytes,
)


class NetworkBlocker(QWebEngineUrlRequestInterceptor):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.blocked_urls: list[str] = []

    def interceptRequest(self, info):
        url = info.requestUrl()
        if url.scheme().lower() in {"http", "https"}:
            self.blocked_urls.append(url.toString())
            info.block(True)


class SmokePage(QWebEnginePage):
    def __init__(self, profile, parent=None):
        super().__init__(profile, parent)
        self.console_errors: list[str] = []

    def javaScriptConsoleMessage(self, level, message, line_number, source_id):
        level_value = getattr(level, "value", level)
        error_level = QWebEnginePage.JavaScriptConsoleMessageLevel.ErrorMessageLevel
        error_value = getattr(error_level, "value", error_level)
        if level_value >= error_value:
            self.console_errors.append(
                f"{source_id}:{line_number}: {message}"
            )


class SmokeBridge(QObject):
    ready_signal = Signal(str)
    loaded_signal = Signal(str)
    saved_signal = Signal(str, str, str)
    error_signal = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.session_id = ""
        self.kind = ""
        self.expected = 0
        self.chunks: list[str | None] = []

    @Slot(str)
    def ready(self, mode):
        self.ready_signal.emit(str(mode or ""))

    @Slot(bool)
    def setDirty(self, dirty):
        if dirty:
            self.error_signal.emit("editor became dirty before user input")

    @Slot(str)
    def editorLoaded(self, session_id):
        self.loaded_signal.emit(str(session_id or ""))

    @Slot(str, str, int)
    def beginPayload(self, session_id, kind, total):
        self.session_id = str(session_id or "")
        self.kind = str(kind or "")
        self.expected = int(total or 0)
        self.chunks = [None] * self.expected

    @Slot(str, int, str)
    def appendPayload(self, session_id, index, chunk):
        if str(session_id or "") != self.session_id:
            self.error_signal.emit("save payload session mismatch")
            return
        index = int(index)
        if index < 0 or index >= len(self.chunks):
            self.error_signal.emit("save payload index out of range")
            return
        self.chunks[index] = str(chunk or "")

    @Slot(str)
    def finishPayload(self, session_id):
        if str(session_id or "") != self.session_id:
            self.error_signal.emit("save payload completion session mismatch")
            return
        if not self.chunks or any(chunk is None for chunk in self.chunks):
            self.error_signal.emit("save payload incomplete")
            return
        self.saved_signal.emit(
            self.session_id,
            self.kind,
            "".join(self.chunks),
        )

    @Slot(str)
    def reportError(self, message):
        self.error_signal.emit(str(message or "unknown editor error"))


def _docx_payload() -> bytes:
    from docx import Document

    output = io.BytesIO()
    document = Document()
    document.add_heading("离线编辑器冒烟测试", level=1)
    document.add_paragraph("正文")
    document.save(output)
    return output.getvalue()


class SmokeCoordinator(QObject):
    def __init__(
        self,
        asset_dir: Path,
        app: QApplication,
        shell_budget_seconds: float,
        model_budget_seconds: float,
        docx_path: Path | None,
    ):
        super().__init__()
        self.asset_dir = asset_dir
        self.app = app
        self.modes = ["html", "docx", "sheet"]
        self.index = -1
        self.failure = ""
        self.session_id = ""
        self.mode = ""
        self.save_requested = False
        self.mode_saved = False
        self.page_started_at = 0.0
        self.model_load_started_at = 0.0
        self.current_shell_seconds = 0.0
        self.shell_budget_seconds = float(shell_budget_seconds)
        self.model_budget_seconds = float(model_budget_seconds)
        self.docx_path = docx_path
        self.docx_session = (
            create_edit_session(str(docx_path))[0]
            if docx_path is not None
            else None
        )
        self.metrics: list[tuple[str, float, float]] = []

        self.profile = QWebEngineProfile("cowork-editor-smoke", self)
        self.blocker = NetworkBlocker(self.profile)
        self.profile.setUrlRequestInterceptor(self.blocker)
        self.page = SmokePage(self.profile, self)
        self.view = QWebEngineView()
        self.view.setPage(self.page)
        self.view.resize(1100, 760)
        self.view.show()

        self.bridge = SmokeBridge(self)
        self.channel = QWebChannel(self.page)
        self.channel.registerObject("deliverableEditorBridge", self.bridge)
        self.page.setWebChannel(self.channel)
        self.bridge.ready_signal.connect(self._handle_ready)
        self.bridge.loaded_signal.connect(self._handle_loaded)
        self.bridge.saved_signal.connect(self._handle_saved)
        self.bridge.error_signal.connect(self.fail)
        self.page.loadFinished.connect(self._handle_page_loaded)

        self.watchdog = QTimer(self)
        self.watchdog.setSingleShot(True)
        self.watchdog.timeout.connect(
            lambda: self.fail(f"{self.mode or 'editor'} smoke test timed out")
        )

    def start(self):
        self._next()

    def _next(self):
        self.index += 1
        if self.index >= len(self.modes):
            if self.blocker.blocked_urls:
                self.fail(
                    "editor attempted remote requests: "
                    + ", ".join(self.blocker.blocked_urls[:3])
                )
                return
            if self.page.console_errors:
                self.fail(
                    "JavaScript console errors: "
                    + " | ".join(self.page.console_errors[:3])
                )
                return
            for mode, shell_seconds, model_seconds in self.metrics:
                print(
                    f"{mode}: shell_ready={shell_seconds:.3f}s "
                    f"model_ready={model_seconds:.3f}s"
                )
            self.watchdog.stop()
            self.app.exit(0)
            return
        self.mode = self.modes[self.index]
        self.session_id = f"smoke-{self.mode}"
        self.save_requested = False
        self.mode_saved = False
        self.page_started_at = time.perf_counter()
        self.model_load_started_at = 0.0
        self.current_shell_seconds = 0.0
        if self.mode == "sheet":
            self.view.resize(581, 700)
        else:
            self.view.resize(1100, 760)
        self.watchdog.start(45000)
        target = self.asset_dir / f"{self.mode}.html"
        if not target.is_file():
            self.fail(f"missing editor asset: {target}")
            return
        self.view.setUrl(QUrl.fromLocalFile(str(target.resolve())))

    def _handle_page_loaded(self, ok):
        if not ok:
            self.fail(f"{self.mode} page failed to load")

    def _payload(self):
        if self.mode == "html":
            return "<!doctype html><html><body><h1>标题</h1><p>正文</p></body></html>"
        if self.mode == "docx":
            payload = (
                self.docx_path.read_bytes()
                if self.docx_path is not None
                else _docx_payload()
            )
            return base64.b64encode(payload).decode("ascii")
        rows = [
            [str(row * 10 + column) for column in range(10)]
            for row in range(10_000)
        ]
        rows[0][0] = "00123"
        rows[0][1] = "9007199254740993"
        return json.dumps(
            rows_to_univer_snapshot(
                rows,
                "十万单元格",
            ),
            ensure_ascii=False,
            separators=(",", ":"),
        )

    def _handle_ready(self, mode):
        if mode != self.mode:
            self.fail(f"ready mode mismatch: expected {self.mode}, got {mode}")
            return
        shell_seconds = time.perf_counter() - self.page_started_at
        if shell_seconds > self.shell_budget_seconds:
            self.fail(
                f"{mode} shell took {shell_seconds:.3f}s, "
                f"exceeding {self.shell_budget_seconds:.3f}s"
            )
            return
        self.current_shell_seconds = shell_seconds
        value = self._payload()
        chunks = [
            value[index : index + 192 * 1024]
            for index in range(0, len(value), 192 * 1024)
        ] or [""]
        self.model_load_started_at = time.perf_counter()
        self.page.runJavaScript(
            "window.coworkEditor.beginLoad("
            + json.dumps(self.session_id)
            + f",{len(chunks)},"
            + json.dumps(self.mode)
            + ")"
        )
        for index, chunk in enumerate(chunks):
            self.page.runJavaScript(
                f"window.coworkEditor.appendLoadChunk({index},"
                + json.dumps(chunk, ensure_ascii=True)
                + ")"
            )
        self.page.runJavaScript("window.coworkEditor.finishLoad()")

    def _handle_loaded(self, session_id):
        if self.save_requested:
            return
        if session_id != self.session_id:
            self.fail("editorLoaded session mismatch")
            return
        model_seconds = time.perf_counter() - self.model_load_started_at
        if model_seconds > self.model_budget_seconds:
            self.fail(
                f"{self.mode} model took {model_seconds:.3f}s, "
                f"exceeding {self.model_budget_seconds:.3f}s"
            )
            return
        self.metrics.append((self.mode, self.current_shell_seconds, model_seconds))
        self.save_requested = True
        if self.mode == "sheet":
            QTimer.singleShot(250, self._probe_sheet_layout)
            return
        QTimer.singleShot(100, self._probe_editor_chrome)

    def _probe_editor_chrome(self):
        self.page.runJavaScript(
            """
            (() => {
              const toolbar = document.querySelector(".editor-toolbar");
              const control = toolbar && toolbar.querySelector("button, select, input");
              const content = document.getElementById("docx-canvas")
                || document.getElementById("html-frame");
              const toolbarRect = toolbar ? toolbar.getBoundingClientRect() : null;
              const controlRect = control ? control.getBoundingClientRect() : null;
              const contentRect = content ? content.getBoundingClientRect() : null;
              return JSON.stringify({
                fontSize: parseFloat(getComputedStyle(document.body).fontSize),
                toolbarHeight: toolbarRect ? toolbarRect.height : 0,
                controlHeight: controlRect ? controlRect.height : 0,
                contentWidth: contentRect ? contentRect.width : 0,
                contentHeight: contentRect ? contentRect.height : 0
              });
            })()
            """,
            self._handle_editor_chrome_probe,
        )

    def _handle_editor_chrome_probe(self, payload):
        try:
            result = json.loads(str(payload or ""))
        except (TypeError, ValueError) as exc:
            self.fail(f"{self.mode} chrome probe returned invalid data: {exc}")
            return
        if float(result.get("fontSize") or 0) < 14:
            self.fail(f"{self.mode} editor font is too small: {result}")
            return
        if float(result.get("toolbarHeight") or 0) < 44:
            self.fail(f"{self.mode} editor toolbar collapsed: {result}")
            return
        if float(result.get("controlHeight") or 0) < 32:
            self.fail(f"{self.mode} editor controls are too small: {result}")
            return
        if float(result.get("contentWidth") or 0) < 900 or float(
            result.get("contentHeight") or 0
        ) < 600:
            self.fail(f"{self.mode} editor content did not fill the viewport: {result}")
            return
        self.page.runJavaScript("window.coworkEditor.requestSave()")

    def _probe_sheet_layout(self):
        self.page.runJavaScript(
            """
            (() => {
              const root = document.getElementById("sheet-root");
              const stylesheet = document.getElementById("sheet-editor-styles");
              const rootRect = root ? root.getBoundingClientRect() : null;
              const canvasRects = [...document.querySelectorAll("#sheet-root canvas")]
                .map((canvas) => canvas.getBoundingClientRect())
                .map((rect) => ({ width: rect.width, height: rect.height }));
              return JSON.stringify({
                stylesheetLoaded: Boolean(
                  stylesheet && stylesheet.sheet && stylesheet.sheet.cssRules.length
                ),
                rootWidth: rootRect ? rootRect.width : 0,
                rootHeight: rootRect ? rootRect.height : 0,
                unresolvedLocale: document.body.innerText.includes("sheets-ui.info."),
                canvasRects
              });
            })()
            """,
            self._handle_sheet_layout_probe,
        )

    def _handle_sheet_layout_probe(self, payload):
        try:
            result = json.loads(str(payload or ""))
        except (TypeError, ValueError) as exc:
            self.fail(f"sheet layout probe returned invalid data: {exc}")
            return
        if not result.get("stylesheetLoaded"):
            self.fail("sheet editor stylesheet was not loaded")
            return
        if float(result.get("rootWidth") or 0) < 500:
            self.fail("sheet editor root did not fill the narrow viewport")
            return
        if float(result.get("rootHeight") or 0) < 600:
            self.fail("sheet editor root height collapsed in the narrow viewport")
            return
        if result.get("unresolvedLocale"):
            self.fail("sheet editor displayed unresolved locale keys")
            return
        canvas_rects = result.get("canvasRects") or []
        if not canvas_rects:
            self.fail(f"sheet grid canvas is not visible: {result}")
            return
        self.page.runJavaScript("window.coworkEditor.requestSave()")

    def _handle_saved(self, session_id, kind, payload):
        if self.mode_saved:
            return
        if session_id != self.session_id or kind != self.mode:
            self.fail("saved payload identity mismatch")
            return
        self.mode_saved = True
        try:
            if kind == "docx":
                docx_bytes = base64.b64decode(payload, validate=True)
                if self.docx_session is not None:
                    docx_bytes = serialize_editor_payload(
                        self.docx_session,
                        docx_bytes,
                    )
                validate_docx_bytes(docx_bytes)
            elif kind == "sheet":
                snapshot = json.loads(payload)
                if not isinstance(snapshot.get("sheets"), dict):
                    raise ValueError("sheet snapshot has no sheets")
                sheet_id = snapshot["sheetOrder"][0]
                first_row = snapshot["sheets"][sheet_id]["cellData"]["0"]
                if first_row["0"].get("v") != "00123":
                    raise ValueError("sheet output coerced a leading-zero string")
                if first_row["1"].get("v") != "9007199254740993":
                    raise ValueError("sheet output coerced a long numeric string")
            elif "<h1" not in payload.lower():
                raise ValueError("HTML output lost edited document body")
        except Exception as exc:
            self.fail(f"{kind} output validation failed: {exc}")
            return
        self.page.runJavaScript(
            "if (window.coworkEditor) window.coworkEditor.dispose()"
        )
        QTimer.singleShot(50, self._next)

    @Slot(str)
    def fail(self, message):
        if self.failure:
            return
        self.failure = str(message or "unknown smoke failure")
        print(f"[ERROR] {self.failure}", file=sys.stderr)
        self.watchdog.stop()
        self.app.exit(1)


def main(argv=None):
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--asset-dir",
        type=Path,
        default=ROOT / "web" / "editors" / "dist",
    )
    parser.add_argument("--shell-budget-seconds", type=float, default=2.5)
    parser.add_argument("--model-budget-seconds", type=float, default=8.0)
    parser.add_argument("--docx-path", type=Path)
    args = parser.parse_args(argv)
    if args.docx_path is not None and not args.docx_path.is_file():
        parser.error(f"DOCX benchmark file does not exist: {args.docx_path}")
    app = QApplication.instance() or QApplication([])
    coordinator = SmokeCoordinator(
        args.asset_dir.resolve(),
        app,
        args.shell_budget_seconds,
        args.model_budget_seconds,
        args.docx_path.resolve() if args.docx_path is not None else None,
    )
    QTimer.singleShot(0, coordinator.start)
    return app.exec()


if __name__ == "__main__":
    raise SystemExit(main())
