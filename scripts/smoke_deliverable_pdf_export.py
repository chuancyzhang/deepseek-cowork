"""End-to-end smoke test for local Markdown, HTML, and DOCX PDF export."""

from __future__ import annotations

import os
from pathlib import Path
import sys

from PySide6.QtCore import QEventLoop, Qt
from PySide6.QtWidgets import QApplication


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from core.deliverable_pdf_export import DeliverablePdfExportController
from main import MainWindow


def _create_inputs(output_dir: Path) -> list[tuple[Path, Path]]:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from PIL import Image, ImageDraw

    output_dir.mkdir(parents=True, exist_ok=True)
    image_path = output_dir / "smoke-local-image.png"
    image = Image.new("RGB", (960, 320), "#f2f0ff")
    draw = ImageDraw.Draw(image)
    draw.rectangle((16, 16, 944, 304), outline="#5b5bd6", width=8)
    draw.text((48, 130), "PDF export local image", fill="#17171a")
    image.save(image_path)

    markdown_path = output_dir / "smoke-markdown.md"
    markdown_path.write_text(
        """# Markdown PDF 导出验证

这是一段用于检查中文字体、分页和可搜索文字的内容。

| 项目 | 状态 |
| --- | --- |
| 中文 | 正常 |
| 表格 | 正常 |

```python
print("Markdown code block")
```

![本地图片](smoke-local-image.png)

## 第二部分

"""
        + "长文本分页检查。" * 900,
        encoding="utf-8",
    )

    html_path = output_dir / "smoke-html.html"
    html_path.write_text(
        """<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8"><style>
body { margin: 18mm; color: #17171a; font: 15px/1.65 "Microsoft YaHei UI", sans-serif; }
h1 { color: #4f46c8; } img { width: 100%; max-width: 720px; }
@media print { .screen-only { display: none; } }
</style></head><body>
<h1>HTML PDF 导出验证</h1><p class="screen-only">这段内容不应出现在 PDF 中。</p>
<p>保留原始 HTML 样式、本地图片和打印媒体规则。</p>
<img src="smoke-local-image.png" alt="本地图片">
<script>document.body.insertAdjacentHTML('beforeend', '<p>JavaScript 渲染完成。</p>');</script>
</body></html>""",
        encoding="utf-8",
    )

    docx_path = output_dir / "smoke-docx.docx"
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    document.add_heading("DOCX PDF 导出验证", 0)
    document.add_paragraph("检查中文、横向纸张、表格和自动分页。")
    table = document.add_table(rows=3, cols=3)
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.text = f"单元格 {row_index + 1}-{column_index + 1}"
    for index in range(90):
        document.add_paragraph(f"分页压力测试段落 {index + 1}：确保长文档生成多个横向页面。")
    document.add_heading("第二页", level=1)
    document.add_paragraph("DOCX 页面应保持横向且没有裁切。")
    document.save(docx_path)

    return [
        (markdown_path, output_dir / "smoke-markdown.pdf"),
        (html_path, output_dir / "smoke-html.pdf"),
        (docx_path, output_dir / "smoke-docx.pdf"),
    ]


def _export(app: QApplication, source: Path, target: Path) -> dict:
    loop = QEventLoop()
    outcome = {}
    controller = DeliverablePdfExportController(app)
    controller.stage_changed.connect(
        lambda stage, details: print(f"{source.name}: {stage} {details}", flush=True)
    )

    def handle_success(result):
        outcome.update(result)
        loop.quit()

    def handle_failure(code, message):
        outcome.update({"error": code, "message": message})
        loop.quit()

    controller.succeeded.connect(handle_success)
    controller.failed.connect(handle_failure)
    controller.start(
        str(source),
        str(target),
        markdown_renderer=lambda text: MainWindow._build_deliverable_markdown_pdf_document(
            None, text
        ),
        docx_editor_path=str(ROOT / "web" / "editors" / "dist" / "docx.html"),
    )
    loop.exec()
    if outcome.get("error"):
        raise RuntimeError(f"{source.name}: {outcome['error']}: {outcome['message']}")
    return outcome


def main() -> int:
    QApplication.setAttribute(Qt.ApplicationAttribute.AA_ShareOpenGLContexts)
    app = QApplication.instance() or QApplication([])
    output_dir = Path(
        os.environ.get("COWORK_PDF_SMOKE_OUTPUT_DIR")
        or ROOT / "tmp" / "pdfs"
    ).resolve()
    results = []
    for source, target in _create_inputs(output_dir):
        results.append(_export(app, source, target))
    for result in results:
        print(
            f"OK {result['target_path']} pages={result['page_count']} "
            f"bytes={result['bytes_written']}",
            flush=True,
        )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
