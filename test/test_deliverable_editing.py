import base64
import io
import json
import os
import shutil
import tempfile
import unittest
import zipfile
from copy import deepcopy
from pathlib import Path
from unittest.mock import patch

from core.deliverable_editing import (
    EDITABLE_EXTENSIONS,
    TEXT_MAX_BYTES,
    DeliverableEditError,
    ExternalModificationError,
    apply_html_dom_patch,
    atomic_save_session,
    create_edit_session,
    editor_descriptor,
    load_editor_payload,
    preflight_edit,
    prepare_html_edit_source,
    restore_html_edit_source,
    restore_previous_version,
    serialize_editor_payload,
    sha256_file,
)


class DeliverableEditingCoreTest(unittest.TestCase):
    def test_registry_covers_planned_formats_only(self):
        for extension in (
            ".docx",
            ".html",
            ".htm",
            ".xlsx",
            ".csv",
            ".tsv",
            ".txt",
            ".md",
            ".markdown",
            ".json",
            ".xml",
            ".yaml",
            ".yml",
            ".log",
        ):
            self.assertIn(extension, EDITABLE_EXTENSIONS)
            self.assertIsNotNone(editor_descriptor(extension))
        for extension in (".pdf", ".pptx", ".doc", ".xls", ".ppt", ".png"):
            self.assertIsNone(editor_descriptor(extension))

    def test_legacy_text_encoding_requires_explicit_choice(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "中文.txt")
            Path(path).write_bytes("中文内容".encode("gb18030"))

            blocked = preflight_edit(path)
            allowed = preflight_edit(path, selected_encoding="gb18030")

        self.assertFalse(blocked.allowed)
        self.assertEqual(blocked.blocking_issues[0].code, "encoding_required")
        self.assertTrue(allowed.allowed)
        self.assertEqual(allowed.encoding, "gb18030")

    def test_structured_text_is_validated_before_serialization(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "data.json")
            Path(path).write_text('{"ready": true}', encoding="utf-8")
            session, _report = create_edit_session(path)

            with self.assertRaisesRegex(DeliverableEditError, "JSON 语法验证失败"):
                serialize_editor_payload(session, '{"ready":')
            with self.assertRaisesRegex(DeliverableEditError, "重复键"):
                serialize_editor_payload(session, '{"ready": true, "ready": false}')
            with self.assertRaisesRegex(DeliverableEditError, "不允许常量"):
                serialize_editor_payload(session, '{"value": NaN}')

    def test_html_edit_copy_preserves_scripts_and_embeds(self):
        source = (
            "<!doctype html><html><head><script>window.secret = 1;</script></head>"
            "<body><h1>旧标题</h1><iframe src='https://example.com'></iframe></body></html>"
        )

        editable, preserved = prepare_html_edit_source(source)
        edited = editable.replace("旧标题", "新标题")
        restored = restore_html_edit_source(edited, preserved)

        self.assertNotIn("<script", editable)
        self.assertNotIn("<iframe", editable)
        self.assertIn("新标题", restored)
        self.assertIn("window.secret = 1", restored)
        self.assertIn("<iframe", restored)

    def test_html_temporary_edit_base_is_not_saved(self):
        source = "<!doctype html><html><head></head><body><img src='image.png'></body></html>"
        editable, preserved = prepare_html_edit_source(
            source,
            "file:///D:/workspace/",
        )

        self.assertIn("data-cowork-edit-base", editable)
        restored = restore_html_edit_source(editable, preserved)

        self.assertNotIn("data-cowork-edit-base", restored)
        self.assertNotIn("file:///D:/workspace/", restored)
        self.assertIn('src="image.png"', restored)

    def test_html_dom_patch_changes_body_but_preserves_original_head(self):
        source = (
            "<!doctype html><html><head><title>原始标题</title>"
            "<style>.brand{color:#123456}</style></head>"
            "<body class='page'><h1>旧正文</h1></body></html>"
        )
        editable, preserved = prepare_html_edit_source(source)
        edited = editable.replace("原始标题", "不应覆盖").replace("旧正文", "新正文")

        patched = apply_html_dom_patch(source, edited, preserved)

        self.assertIn("<title>原始标题</title>", patched)
        self.assertNotIn("不应覆盖", patched)
        self.assertIn(".brand{color:#123456}", patched)
        self.assertIn("新正文", patched)
        self.assertIn('class="page"', patched)

    def test_html_reserved_editor_attributes_block_entry(self):
        with self.assertRaisesRegex(DeliverableEditError, "保留"):
            prepare_html_edit_source(
                "<html><body data-cowork-edit-base='true'>正文</body></html>"
            )

    def test_missing_html_safety_placeholder_blocks_save(self):
        editable, preserved = prepare_html_edit_source(
            "<html><body><script>safe()</script><p>content</p></body></html>"
        )
        edited = editable.replace(
            '<template data-cowork-preserved-node="0"></template>',
            "",
        )

        with self.assertRaisesRegex(DeliverableEditError, "缺少受保护"):
            restore_html_edit_source(edited, preserved)

    def test_docx_common_document_is_editable_and_tracked_changes_are_blocked(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "report.docx")
            document = Document()
            document.add_heading("报告", level=1)
            document.add_paragraph("正文")
            document.save(path)
            self.assertTrue(preflight_edit(path).allowed)

            changed = os.path.join(directory, "tracked.docx")
            with zipfile.ZipFile(path) as source, zipfile.ZipFile(changed, "w") as target:
                for item in source.infolist():
                    data = source.read(item.filename)
                    if item.filename == "word/document.xml":
                        data = data.replace(b"<w:body>", b"<w:body><w:ins/>", 1)
                    target.writestr(item, data)

            before_preflight = Path(changed).read_bytes()
            report = preflight_edit(changed)
            after_preflight = Path(changed).read_bytes()

        self.assertFalse(report.allowed)
        self.assertTrue(any(issue.code == "docx_unsupported_markup" for issue in report.issues))
        self.assertEqual(after_preflight, before_preflight)

    def test_docx_plain_header_and_footer_are_preserved_while_body_changes(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "report.docx")
            document = Document()
            document.add_paragraph("原正文")
            document.sections[0].header.paragraphs[0].text = "固定页眉"
            document.sections[0].footer.paragraphs[0].text = "固定页脚"
            document.save(path)

            session, report = create_edit_session(path)
            self.assertTrue(report.allowed)
            self.assertEqual(
                report.metadata.get("docx_preserved_header_footer_count"),
                2,
            )

            edited_stream = io.BytesIO()
            edited = Document()
            edited.add_paragraph("编辑后的正文")
            edited.save(edited_stream)
            output = serialize_editor_payload(session, edited_stream.getvalue())

            round_trip = Document(io.BytesIO(output))
            self.assertEqual(
                [paragraph.text for paragraph in round_trip.paragraphs],
                ["编辑后的正文"],
            )
            self.assertEqual(
                round_trip.sections[0].header.paragraphs[0].text,
                "固定页眉",
            )
            self.assertEqual(
                round_trip.sections[0].footer.paragraphs[0].text,
                "固定页脚",
            )
            with zipfile.ZipFile(io.BytesIO(output)) as archive:
                names = set(archive.namelist())
                self.assertTrue(
                    any(name.startswith("word/cowork-preserved-header") for name in names)
                )
                self.assertTrue(
                    any(name.startswith("word/cowork-preserved-footer") for name in names)
                )

    def test_docx_header_with_related_image_remains_blocked(self):
        from docx import Document
        from docx.shared import Inches

        image_bytes = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "image-header.docx")
            document = Document()
            document.add_paragraph("正文")
            document.sections[0].header.paragraphs[0].add_run().add_picture(
                io.BytesIO(image_bytes),
                width=Inches(0.1),
            )
            document.save(path)

            report = preflight_edit(path)

        self.assertFalse(report.allowed)
        self.assertTrue(
            any(issue.code == "docx_complex_header_footer" for issue in report.issues)
        )

    def test_xlsx_snapshot_round_trip_keeps_values_formula_style_and_merge(self):
        import openpyxl
        from openpyxl.styles import Font, PatternFill

        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "data.xlsx")
            workbook = openpyxl.Workbook()
            worksheet = workbook.active
            worksheet.title = "明细"
            worksheet["A1"] = "标题"
            worksheet["A1"].font = Font(bold=True)
            worksheet["A1"].fill = PatternFill(fill_type="solid", fgColor="FFFF00")
            worksheet["B2"] = "=1+2"
            worksheet.merge_cells("A3:B3")
            worksheet.freeze_panes = "B2"
            worksheet.sheet_properties.tabColor = "FF336699"
            workbook.save(path)
            workbook.close()

            session, _report = create_edit_session(path)
            snapshot = load_editor_payload(session)["snapshot"]
            output = serialize_editor_payload(session, snapshot)
            output_path = os.path.join(directory, "output.xlsx")
            Path(output_path).write_bytes(output)
            round_trip = openpyxl.load_workbook(output_path, data_only=False)
            try:
                worksheet = round_trip["明细"]
                self.assertEqual(worksheet["A1"].value, "标题")
                self.assertTrue(worksheet["A1"].font.bold)
                self.assertEqual(worksheet["B2"].value, "=1+2")
                self.assertIn("A3:B3", {str(item) for item in worksheet.merged_cells.ranges})
                self.assertEqual(str(worksheet.freeze_panes), "B2")
                self.assertEqual(worksheet.sheet_properties.tabColor.rgb, "FF336699")
            finally:
                round_trip.close()

    def test_xlsx_with_chart_is_blocked(self):
        import openpyxl
        from openpyxl.chart import BarChart, Reference

        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "chart.xlsx")
            workbook = openpyxl.Workbook()
            worksheet = workbook.active
            worksheet.append(["名称", "数值"])
            worksheet.append(["A", 1])
            chart = BarChart()
            chart.add_data(Reference(worksheet, min_col=2, min_row=1, max_row=2), titles_from_data=True)
            worksheet.add_chart(chart, "D2")
            workbook.save(path)
            workbook.close()

            report = preflight_edit(path)

        self.assertFalse(report.allowed)
        self.assertTrue(any(issue.code == "xlsx_unsupported_part" for issue in report.issues))

    def test_atomic_save_keeps_one_backup_detects_conflict_and_restores_by_swap(self):
        with tempfile.TemporaryDirectory() as directory:
            backup_root = os.path.join(directory, "backups")
            path = os.path.join(directory, "note.txt")
            Path(path).write_text("one\n", encoding="utf-8")
            session, _report = create_edit_session(path)

            first = atomic_save_session(
                session,
                serialize_editor_payload(session, "two\n"),
                backup_root=backup_root,
            )
            self.assertEqual(Path(first.backup_path).read_text(encoding="utf-8"), "one\n")
            atomic_save_session(
                session,
                serialize_editor_payload(session, "three\n"),
                backup_root=backup_root,
            )
            self.assertEqual(Path(first.backup_path).read_text(encoding="utf-8"), "two\n")

            restored = restore_previous_version(path, backup_root=backup_root)
            self.assertEqual(Path(path).read_text(encoding="utf-8"), "two\n")
            self.assertEqual(Path(restored.backup_path).read_text(encoding="utf-8"), "three\n")

            conflict_session, _report = create_edit_session(path)
            Path(path).write_text("external\n", encoding="utf-8")
            before = sha256_file(path)
            with self.assertRaises(ExternalModificationError):
                atomic_save_session(
                    conflict_session,
                    serialize_editor_payload(conflict_session, "mine\n"),
                    backup_root=backup_root,
                )
            self.assertEqual(sha256_file(path), before)
            self.assertEqual(Path(path).read_text(encoding="utf-8"), "external\n")

    def test_over_limit_and_deleted_files_are_blocked_without_writes(self):
        with tempfile.TemporaryDirectory() as directory:
            oversized = os.path.join(directory, "large.txt")
            with open(oversized, "wb") as handle:
                handle.truncate(TEXT_MAX_BYTES + 1)
            before = os.path.getsize(oversized)

            report = preflight_edit(oversized)

            self.assertFalse(report.allowed)
            self.assertEqual(report.blocking_issues[0].code, "file_too_large")
            self.assertEqual(os.path.getsize(oversized), before)

            source = os.path.join(directory, "deleted.txt")
            Path(source).write_text("original", encoding="utf-8")
            session, _report = create_edit_session(source)
            data = serialize_editor_payload(session, "edited")
            os.unlink(source)

            with self.assertRaisesRegex(DeliverableEditError, "原文件已被删除"):
                atomic_save_session(session, data, backup_root=os.path.join(directory, "backups"))
            self.assertFalse(os.path.exists(source))

    def test_atomic_save_failure_leaves_original_bytes_unchanged(self):
        with tempfile.TemporaryDirectory() as directory:
            source = os.path.join(directory, "failure.txt")
            original = b"original\r\n"
            Path(source).write_bytes(original)
            session, _report = create_edit_session(source)
            data = serialize_editor_payload(session, "edited\n")

            with patch(
                "core.deliverable_editing.os.replace",
                side_effect=OSError("disk write denied"),
            ):
                with self.assertRaisesRegex(DeliverableEditError, "保存文件失败"):
                    atomic_save_session(
                        session,
                        data,
                        backup_root=os.path.join(directory, "backups"),
                    )

            self.assertEqual(Path(source).read_bytes(), original)

    def test_atomic_save_rechecks_source_after_backup_copy(self):
        with tempfile.TemporaryDirectory() as directory:
            source = os.path.join(directory, "race.txt")
            Path(source).write_text("original", encoding="utf-8")
            session, _report = create_edit_session(source)
            data = serialize_editor_payload(session, "mine")
            original_copy = shutil.copy2

            def copy_then_external_write(src, dst, *args, **kwargs):
                result = original_copy(src, dst, *args, **kwargs)
                if os.path.normcase(os.path.abspath(src)) == os.path.normcase(source):
                    Path(source).write_text("external", encoding="utf-8")
                return result

            with patch(
                "core.deliverable_editing.shutil.copy2",
                side_effect=copy_then_external_write,
            ):
                with self.assertRaises(ExternalModificationError):
                    atomic_save_session(
                        session,
                        data,
                        backup_root=os.path.join(directory, "backups"),
                    )

            self.assertEqual(Path(source).read_text(encoding="utf-8"), "external")

    def test_csv_uses_sheet_snapshot_and_preserves_quoted_values(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "data.csv")
            Path(path).write_text('name,note\r\nAlice,"a,b"\r\n', encoding="utf-8", newline="")
            session, _report = create_edit_session(path)
            snapshot = load_editor_payload(session)["snapshot"]

            output = serialize_editor_payload(session, snapshot).decode("utf-8")

        self.assertIn('"a,b"', output)
        self.assertIn("\r\n", output)

    def test_csv_rejects_extra_sheets_and_sheet_indices_outside_xlsx_bounds(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "data.csv")
            Path(path).write_text("name,value\nA,1\n", encoding="utf-8")
            session, _report = create_edit_session(path)
            snapshot = load_editor_payload(session)["snapshot"]

            extra_sheet = deepcopy(next(iter(snapshot["sheets"].values())))
            extra_sheet["id"] = "extra"
            extra_sheet["name"] = "额外"
            snapshot["sheets"]["extra"] = extra_sheet
            snapshot["sheetOrder"].append("extra")
            with self.assertRaisesRegex(DeliverableEditError, "只能保存一个工作表"):
                serialize_editor_payload(session, snapshot)

            snapshot["sheetOrder"] = snapshot["sheetOrder"][:1]
            snapshot["sheets"].pop("extra")
            sheet = snapshot["sheets"][snapshot["sheetOrder"][0]]
            sheet["cellData"]["1048576"] = {"0": {"v": "越界", "t": 1}}
            with self.assertRaisesRegex(DeliverableEditError, "行索引超出"):
                serialize_editor_payload(session, snapshot)


if __name__ == "__main__":
    unittest.main()
