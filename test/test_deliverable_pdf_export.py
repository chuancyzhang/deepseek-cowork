import io
import os
from pathlib import Path
import tempfile
import unittest
from unittest.mock import MagicMock, patch

from PySide6.QtGui import QAction
from PySide6.QtWidgets import QApplication
from pypdf import PdfReader, PdfWriter

from core.deliverable_pdf_export import (
    DeliverablePdfExportController,
    DeliverablePdfExportError,
    _prepare_source,
    default_pdf_target,
    validate_pdf_bytes,
    write_pdf_atomic,
)
from main import MainWindow


ROOT = Path(__file__).resolve().parents[1]


def _one_page_pdf_bytes():
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    stream = io.BytesIO()
    writer.write(stream)
    return stream.getvalue()


class DeliverablePdfExportCoreTest(unittest.TestCase):
    def test_default_target_replaces_supported_source_extension(self):
        self.assertEqual(
            default_pdf_target(os.path.join("C:\\workspace", "report.markdown")),
            os.path.abspath(os.path.join("C:\\workspace", "report.pdf")),
        )

    def test_validate_and_atomic_write_produce_reopenable_pdf(self):
        data = _one_page_pdf_bytes()
        self.assertEqual(validate_pdf_bytes(data), 1)
        with tempfile.TemporaryDirectory() as directory:
            target = os.path.join(directory, "report.pdf")
            self.assertEqual(write_pdf_atomic(target, data), 1)
            self.assertEqual(len(PdfReader(target).pages), 1)
            self.assertFalse(any(name.endswith(".tmp") for name in os.listdir(directory)))

    def test_invalid_pdf_never_replaces_existing_target(self):
        with tempfile.TemporaryDirectory() as directory:
            target = os.path.join(directory, "report.pdf")
            Path(target).write_bytes(b"existing")
            with self.assertRaises(DeliverablePdfExportError) as raised:
                write_pdf_atomic(target, b"not-a-pdf")
            self.assertEqual(raised.exception.code, "pdf_invalid")
            self.assertEqual(Path(target).read_bytes(), b"existing")

    def test_source_removed_during_render_never_writes_target(self):
        app = QApplication.instance() or QApplication([])
        with tempfile.TemporaryDirectory() as directory:
            source = os.path.join(directory, "report.md")
            target = os.path.join(directory, "report.pdf")
            Path(source).write_text("# Draft", encoding="utf-8")
            controller = DeliverablePdfExportController()
            controller._source_path = source
            controller._target_path = target
            controller._started_at = 1.0
            failures = []
            controller.failed.connect(lambda code, message: failures.append((code, message)))
            os.remove(source)

            controller._handle_pdf_bytes(_one_page_pdf_bytes())

            self.assertEqual(failures[0][0], "source_removed_during_export")
            self.assertFalse(os.path.exists(target))
        self.assertIsNotNone(app)

    def test_prepare_markdown_requires_utf8_and_supported_format(self):
        with tempfile.TemporaryDirectory() as directory:
            markdown_path = os.path.join(directory, "report.md")
            Path(markdown_path).write_text("# 中文\n\n内容", encoding="utf-8")
            prepared = _prepare_source(markdown_path)
            self.assertEqual(prepared["kind"], "markdown")
            self.assertIn("中文", prepared["text"])

            invalid_path = os.path.join(directory, "legacy.md")
            Path(invalid_path).write_bytes("中文".encode("gbk"))
            with self.assertRaises(DeliverablePdfExportError) as raised:
                _prepare_source(invalid_path)
            self.assertEqual(raised.exception.code, "markdown_encoding_invalid")

            text_path = os.path.join(directory, "notes.txt")
            Path(text_path).write_text("notes", encoding="utf-8")
            with self.assertRaises(DeliverablePdfExportError) as raised:
                _prepare_source(text_path)
            self.assertEqual(raised.exception.code, "source_format_unsupported")

    def test_docx_manual_page_break_is_rejected_instead_of_dropped(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as directory:
            source = os.path.join(directory, "manual-break.docx")
            document = Document()
            document.add_paragraph("第一页")
            document.add_page_break()
            document.add_paragraph("第二页")
            document.save(source)
            with self.assertRaises(DeliverablePdfExportError) as raised:
                _prepare_source(source)
            self.assertEqual(
                raised.exception.code,
                "docx_manual_page_break_unsupported",
            )

    def test_docx_inline_image_is_rejected_instead_of_dropped(self):
        from docx import Document
        from docx.shared import Inches
        from PIL import Image

        with tempfile.TemporaryDirectory() as directory:
            image_path = os.path.join(directory, "image.png")
            Image.new("RGB", (32, 32), "red").save(image_path)
            source = os.path.join(directory, "with-image.docx")
            document = Document()
            document.add_picture(image_path, width=Inches(1))
            document.save(source)
            with self.assertRaises(DeliverablePdfExportError) as raised:
                _prepare_source(source)
            self.assertEqual(
                raised.exception.code,
                "docx_inline_image_unsupported",
            )


class DeliverablePdfExportUiTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.app = QApplication.instance() or QApplication([])

    def test_export_action_is_visible_only_for_requested_formats(self):
        with tempfile.TemporaryDirectory() as directory:
            window = MainWindow.__new__(MainWindow)
            window.deliverable_more_export_pdf_action = QAction("导出为 PDF…")
            for extension, expected in (
                (".md", True),
                (".markdown", True),
                (".html", True),
                (".htm", True),
                (".docx", True),
                (".txt", False),
                (".pdf", False),
            ):
                path = os.path.join(directory, f"sample{extension}")
                Path(path).write_bytes(b"sample")
                window.current_deliverable_path = path
                window._sync_deliverable_action_visibility()
                self.assertEqual(
                    window.deliverable_more_export_pdf_action.isVisible(),
                    expected,
                    msg=extension,
                )

    def test_dirty_file_is_saved_before_export_prompt(self):
        with tempfile.TemporaryDirectory() as directory:
            source = os.path.join(directory, "report.md")
            Path(source).write_text("# Draft", encoding="utf-8")
            window = MainWindow.__new__(MainWindow)
            window.current_deliverable_path = source
            window.deliverable_pdf_export_controller = None
            window.deliverable_edit_state = "dirty"
            window.deliverable_edit_dirty = True
            window.deliverable_edit_session = type("Session", (), {"path": source})()
            window.deliverable_pending_pdf_export_source = ""
            window.save_deliverable_edit = MagicMock()
            window._prompt_and_start_deliverable_pdf_export = MagicMock()
            with patch("main.log_sub_agent_runtime"), patch(
                "main.ProductMessageDialog"
            ) as dialog_type:
                dialog_type.return_value.exec_result.return_value = "save_export"
                window.export_current_deliverable_pdf()

            self.assertEqual(window.deliverable_pending_pdf_export_source, os.path.abspath(source))
            window.save_deliverable_edit.assert_called_once_with()
            window._prompt_and_start_deliverable_pdf_export.assert_not_called()

    def test_legacy_html_generate_pdf_action_is_removed(self):
        source = (ROOT / "main.py").read_text(encoding="utf-8")
        self.assertNotIn('generate_pdf_action = QAction("生成 PDF"', source)
        self.assertIn('generate_docx_action = QAction("生成 DOCX"', source)
        self.assertIn('"导出为 PDF…"', source)


if __name__ == "__main__":
    unittest.main()
